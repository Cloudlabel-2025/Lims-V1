from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_lims_role_flow_document import (
    add_body, add_callout, add_figure, add_heading, add_page_number,
    add_table, font, hex_rgb, rounded_box, draw_centered_text, arrow,
    set_run, NAVY, BLUE, GRAY, BLACK, WHITE, LIGHT_BLUE, PALE,
    GREEN, AMBER, RED, RED_DARK, AMBER_DARK,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "deliverables"
ASSETS = ROOT / "docs" / ".qa_split_doc_assets"
TESTER_DOC = OUT / "LIMS_Simple_Testing_Flow_and_Scenarios.docx"
REFERENCE_DOC = OUT / "LIMS_Roles_Gaps_and_Product_Reference.docx"


def configure(doc, header_text, title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.2

    for level, size, before, after, color in [
        (1, 18, 18, 9, NAVY),
        (2, 15, 14, 7, BLUE),
        (3, 13, 10, 5, "1F4D78"),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(12)
        style.paragraph_format.left_indent = Inches(0.42)
        style.paragraph_format.first_line_indent = Inches(-0.20)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.2

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run(header_text)
    set_run(r, size=8.5, bold=True, color=GRAY)
    add_page_number(section.footer.paragraphs[0])
    doc.core_properties.title = title
    doc.core_properties.author = "LIMS Product Team"


def title_page(doc, kicker, title, subtitle, audience):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(38)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(kicker)
    set_run(r, size=11, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    set_run(r, size=27, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(subtitle)
    set_run(r, size=14, color=GRAY)
    add_callout(doc, "For", audience, "info")


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run(r, size=12)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run(r, size=12)


def label_para(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{label}: ")
    set_run(r, size=12, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run(r, size=12)
    return p


def scenario(doc, sid, name, role, start, steps, pass_if, evidence):
    add_heading(doc, f"{sid} - {name}", 2)
    label_para(doc, "Test as", role)
    label_para(doc, "Before you start", start)
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    r = p.add_run("Steps")
    set_run(r, size=12, bold=True, color=NAVY)
    numbered(doc, steps)
    add_callout(doc, "PASS if", pass_if, "ok")
    label_para(doc, "Save as evidence", evidence)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Result:  PASS / FAIL / BLOCKED     Defect ID: __________________")
    set_run(r, size=11, bold=True, color=GRAY)


def save_simple_flow(path):
    w, h = 1600, 2100
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 45), "Test the LIMS in this order", font=font(45, True), fill=hex_rgb(NAVY))
    d.text((70, 108), "Finish one step before moving to the next step.", font=font(27), fill=hex_rgb(GRAY))
    steps = [
        ("1", "ADMIN", "Create users and give each user the correct role", LIGHT_BLUE),
        ("2", "RECEPTION", "Find or register the patient", "EEF2FF"),
        ("3", "RECEPTION / CASHIER", "Add tests, create the bill and record payment", "EEF2FF"),
        ("4", "PHLEBOTOMIST", "Check the patient, collect the sample and label it", GREEN),
        ("5", "LAB TECHNICIAN", "Process the sample and enter the test results", GREEN),
        ("6", "PATHOLOGIST", "Review, approve and release the report", AMBER),
        ("7", "DOCTOR", "Open only referred patients and released reports", LIGHT_BLUE),
        ("8", "PATIENT", "Open own bills and released reports", LIGHT_BLUE),
    ]
    x1, x2, top, bh, gap = 170, 1430, 185, 170, 70
    previous = None
    for i, (num, role, action, fill) in enumerate(steps):
        y1 = top + i * (bh + gap)
        box = (x1, y1, x2, y1 + bh)
        rounded_box(d, box, fill, BLUE, 20, 4)
        d.ellipse((x1 + 20, y1 + 35, x1 + 120, y1 + 135), fill=hex_rgb(NAVY))
        draw_centered_text(d, (x1 + 20, y1 + 35, x1 + 120, y1 + 135), num, 34, True, WHITE, 3)
        draw_centered_text(d, (x1 + 145, y1 + 20, x1 + 505, y1 + 150), role, 24, True, NAVY, 23)
        draw_centered_text(d, (x1 + 525, y1 + 15, x2 - 25, y1 + 155), action, 26, False, BLACK, 42)
        if previous:
            arrow(d, previous, ((x1 + x2)//2, y1 - 5), BLUE, 5, 14)
        previous = ((x1 + x2)//2, y1 + bh + 5)
    im.save(path, quality=95)


def save_handoff_flow(path):
    w, h = 1600, 1150
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    d.text((65, 45), "Who hands the work to whom?", font=font(44, True), fill=hex_rgb(NAVY))
    items = [
        ("Reception", "Patient + bill"),
        ("Phlebotomist", "Collected sample"),
        ("Lab Technician", "Draft results"),
        ("Pathologist", "Released report"),
        ("Doctor / Patient", "View report"),
    ]
    y = 210
    boxes = []
    for i, (role, output) in enumerate(items):
        x1 = 55 + i * 307
        box = (x1, y, x1 + 260, y + 390)
        boxes.append(box)
        rounded_box(d, box, ["EEF2FF", GREEN, GREEN, AMBER, LIGHT_BLUE][i], BLUE, 18, 4)
        draw_centered_text(d, (x1+15, y+25, x1+245, y+145), role, 25, True, NAVY, 18)
        draw_centered_text(d, (x1+20, y+165, x1+240, y+345), output, 27, False, BLACK, 17)
        if i:
            arrow(d, (boxes[i-1][2]+3, y+195), (x1-3, y+195), BLUE, 5, 13)
    d.rounded_rectangle((150, 760, 1450, 955), radius=18, fill=hex_rgb(RED), outline=hex_rgb(RED_DARK), width=4)
    draw_centered_text(d, (175, 775, 1425, 940), "Important: A Lab Technician enters results.\nA Pathologist or Lab Manager releases the report.", 30, True, RED_DARK, 64)
    im.save(path, quality=95)


def build_tester_doc(flow_img, handoff_img):
    doc = Document()
    configure(doc, "LIMS | SIMPLE TESTING HANDBOOK", "LIMS Simple Testing Flow and Scenarios")
    title_page(
        doc,
        "TESTING HANDBOOK",
        "LIMS Simple Testing Flow and Scenarios",
        "A step-by-step guide written for manual testing and UAT",
        "Testers who need to test the complete LIMS flow from patient registration to report viewing.",
    )
    add_heading(doc, "How to use this document", 1)
    numbered(doc, [
        "Create the test users and test data listed below.",
        "Run the main flow in the shown order.",
        "For every scenario, follow the steps exactly.",
        "Compare the screen with the PASS result.",
        "Save the requested evidence.",
        "Mark the scenario as PASS, FAIL or BLOCKED.",
    ])
    add_callout(doc, "Simple rule", "If the actual result is different from the PASS result, create a defect.", "warn")

    doc.add_page_break()
    add_heading(doc, "1. Main testing flow", 1)
    add_figure(doc, flow_img, "Figure 1. Run the main LIMS test in this order.", 6.4)
    add_heading(doc, "What each role does", 2)
    add_figure(doc, handoff_img, "Figure 2. Work hand-off between the main roles.", 6.4)

    doc.add_page_break()
    add_heading(doc, "2. Prepare before testing", 1)
    add_heading(doc, "2.1 Test users", 2)
    add_table(doc, ["User", "Role", "Used for"], [
        ("admin.test", "Admin", "Users, roles and master setup"),
        ("frontdesk.test", "Front Desk Receptionist", "Patient and billing"),
        ("cashier.test", "Billing Cashier", "Payments and refunds"),
        ("phlebo.test", "Phlebotomist", "Sample collection"),
        ("tech.test", "Lab Technician", "Sample processing and result entry"),
        ("path.test", "Pathologist", "Report review and release"),
        ("inventory.test", "Inventory Manager", "Stock"),
        ("doctor.regular", "Doctor Regular", "Referred patient portal"),
        ("doctor.investor", "Doctor Investor", "Investor access"),
        ("Patient account", "Patient portal", "Own bills and reports"),
    ], [2400, 2800, 4160], font_size=9.5)
    add_heading(doc, "2.2 Test data", 2)
    bullets(doc, [
        "One new patient who is not in the system.",
        "One existing patient with previous visits.",
        "One regular doctor and one investor doctor.",
        "One active individual test and one active test package.",
        "One test with numeric results and one test with text results.",
        "One inventory item with enough stock and one item with low stock.",
        "Payment data for cash, card and UPI.",
    ])
    add_heading(doc, "2.3 Evidence naming", 2)
    add_body(doc, "Use this file name: Scenario ID - User - Record ID - Step. Example: T04-frontdesk-BILL-000101-created.png")

    doc.add_page_break()
    add_heading(doc, "3. Admin and access scenarios", 1)
    scenario(doc, "T01", "Create users and assign roles", "Admin", "The required roles are available.", [
        "Open Settings and Users.",
        "Create one user for each test role.",
        "Assign the correct role to each user.",
        "Sign out.",
        "Sign in as each user.",
    ], "Every user can sign in. Each user sees only the menus allowed for that role.", "User list and one screenshot from each role after login.")
    scenario(doc, "T02", "Block an inactive user", "Admin and inactive test user", "Create a test user and confirm that the user can sign in.", [
        "Sign in as Admin.",
        "Make the test user inactive.",
        "Sign out.",
        "Try to sign in as the inactive user.",
    ], "The inactive user cannot sign in or open protected pages.", "Inactive user setting and denied login message.")
    scenario(doc, "T03", "Check role restrictions", "Lab Technician", "The Lab Technician user is active.", [
        "Sign in as Lab Technician.",
        "Try to collect a sample.",
        "Try to release a report.",
        "Try the same actions using the page URL or API request if available.",
    ], "Both actions are denied because the default Lab Technician role does not have these permissions.", "Denied messages or API status codes.")

    doc.add_page_break()
    add_heading(doc, "4. Patient and billing scenarios", 1)
    scenario(doc, "T04", "Register a new patient", "Front Desk Receptionist", "Use a phone number that is not already in the system.", [
        "Open Patients.",
        "Search by phone number.",
        "Confirm that no patient is found.",
        "Select Add Patient.",
        "Enter all required details.",
        "Save the patient.",
    ], "A patient ID is created. The patient appears in search. The saved details are correct.", "Patient ID, patient page and search result.")
    scenario(doc, "T05", "Check a possible duplicate patient", "Front Desk Receptionist", "Use the name and phone number of an existing patient.", [
        "Open Add Patient.",
        "Enter the existing patient's details.",
        "Try to save.",
        "Check whether the system warns or blocks the duplicate.",
    ], "The system follows the agreed duplicate-patient rule. It does not silently create an unsafe duplicate.", "Warning or result, plus the final patient search.")
    scenario(doc, "T06", "Create a bill and test order", "Front Desk Receptionist", "A patient, active doctor, active test and active package exist.", [
        "Open Billing and select the patient.",
        "Select the referral doctor.",
        "Add one individual test and one package.",
        "Select Routine or Urgent priority.",
        "Check the subtotal, discount, tax and total.",
        "Create the bill.",
    ], "A bill ID is created. The total is correct. Linked samples are created. The bill appears in billing history.", "Bill, invoice, total and linked sample IDs.")
    scenario(doc, "T07", "Record partial and final payment", "Billing Cashier", "Use an unpaid bill.", [
        "Open the bill.",
        "Pay less than the full amount.",
        "Check the remaining balance.",
        "Pay the remaining amount.",
        "Open payment history and receipts.",
    ], "The bill changes from Unpaid to Partial and then Paid. The balance becomes zero. Both payments are saved.", "Billing status, balance, receipts and payment history.")
    scenario(doc, "T08", "Block an invalid payment", "Billing Cashier", "Use a bill with a remaining balance.", [
        "Try to pay more than the remaining balance.",
        "Try an unauthorized discount or refund with a user who lacks permission.",
    ], "The system blocks the invalid amount and unauthorized action. No extra receipt or journal is created.", "Error message and unchanged payment history.")

    doc.add_page_break()
    add_heading(doc, "5. Sample scenarios", 1)
    scenario(doc, "T09", "Collect and process a sample", "Phlebotomist, then Lab Technician", "Use a new bill with a Registered sample.", [
        "Sign in as Phlebotomist.",
        "Open the sample worklist.",
        "Check the patient ID, bill ID, test and sample type.",
        "Record the collection and received details.",
        "Mark the sample Collected.",
        "Sign in as Lab Technician.",
        "Move the sample to Processing.",
    ], "The sample moves in order: Registered, Collected, Processing. Times and user names are saved. The bill item is updated.", "Sample status, barcode, custody history and bill item status.")
    scenario(doc, "T10", "Reject a bad sample", "Phlebotomist", "Use a sample that is not Released or Archived.", [
        "Open the sample.",
        "Select Reject.",
        "Try to save without a reason.",
        "Enter a rejection reason.",
        "Save the rejection.",
    ], "The system requires a reason. The sample becomes Rejected. The custody history keeps the rejection event.", "Rejected status, reason and custody history.")
    scenario(doc, "T11", "Check barcode and patient mismatch", "Phlebotomist", "Prepare two patients with two different sample IDs.", [
        "Open Patient A's order.",
        "Scan or enter Patient B's sample ID or barcode.",
        "Try to collect or process the sample.",
    ], "The wrong sample cannot be processed for the patient.", "Mismatch message and unchanged sample records.")

    doc.add_page_break()
    add_heading(doc, "6. Result and report scenarios", 1)
    scenario(doc, "T12", "Enter test results", "Lab Technician or Report Typist", "Use a sample in Processing with configured result fields.", [
        "Open the sample or report entry screen.",
        "Leave one required result empty and try to complete.",
        "Enter an invalid numeric value and try again.",
        "Enter valid results.",
        "Save and complete the sample.",
    ], "Missing or invalid results are blocked. Valid results are saved. A Draft report is available.", "Validation messages, saved values, sample status and Draft report ID.")
    scenario(doc, "T13", "Review and release a report", "Pathologist", "A Draft report with valid results exists.", [
        "Open the Draft report.",
        "Check the patient, sample, test, values, units, ranges and flags.",
        "Move the report to Reviewed.",
        "Move it to Approved.",
        "Move it to Released.",
        "Try to skip one status in a separate test.",
    ], "The normal status order works. Skipping a status is blocked. Reviewer, approver and release details are saved.", "All four statuses and the user/time details.")
    scenario(doc, "T14", "Hide unfinished reports from portals", "Doctor and Patient", "Create one Draft, one Reviewed, one Approved and one Released report.", [
        "Sign in as the linked Doctor.",
        "Check the patient's reports.",
        "Sign in as the Patient.",
        "Check Reports.",
        "Try direct links to the unfinished reports.",
    ], "Only the Released report is visible. Draft, Reviewed and Approved reports are not returned.", "Doctor view, patient view and denied direct-link results.")

    doc.add_page_break()
    add_heading(doc, "7. Doctor and patient portal scenarios", 1)
    scenario(doc, "T15", "Activate and use the doctor portal", "Doctor Regular", "Staff has created the doctor with a unique email.", [
        "Open the activation link.",
        "Enter the email, lab ID and OTP.",
        "Create a password.",
        "Sign in.",
        "Open a referred patient and a released report.",
        "Try to open a patient referred to another doctor.",
    ], "Activation works once. The doctor sees only owned referrals and Released reports. The other doctor's patient is denied.", "Activation result, owned patient, released report and denied access.")
    scenario(doc, "T16", "Activate and use the patient portal", "Patient", "Staff has printed a new patient QR/PIN slip.", [
        "Scan the QR code.",
        "Enter the access PIN and date of birth.",
        "Set the portal PIN.",
        "Open Bills and Reports.",
        "Sign out and sign in with Patient ID, DOB and portal PIN.",
    ], "The patient sees only their own bills, receipts and Released reports. Returning login works.", "Activation, Bills, Reports and returning login.")
    scenario(doc, "T17", "Lock and reset patient portal access", "Patient and Front Desk", "The patient portal account is active.", [
        "Enter the wrong portal PIN five times.",
        "Confirm temporary lockout.",
        "Ask staff to print a new access slip.",
        "Activate with the new slip.",
        "Try the old slip and old session again.",
    ], "Lockout occurs. The new slip works. The old slip and old sessions no longer work.", "Lockout, new activation and denied old credentials.")

    doc.add_page_break()
    add_heading(doc, "8. Inventory, accounts and security scenarios", 1)
    scenario(doc, "T18", "Check stock used by a test", "Inventory Manager and Lab Technician", "Link an inventory item to a test. Prepare enough stock and low stock cases.", [
        "Run the test with enough stock.",
        "Check on-hand and reserved stock.",
        "Run the test with low or expired stock.",
        "Check the warning or block behavior.",
    ], "Stock values follow the implemented rule. Low/expired stock is not silently ignored. Any incomplete automatic consumption is recorded as a known gap.", "Stock before/after, reservation and warning/block.")
    scenario(doc, "T19", "Check bill, receipt and account totals", "Accounts Manager", "Use the bill from T06 and payments from T07.", [
        "Open the invoice, receipts and ledger.",
        "Compare bill total, paid amount and balance.",
        "Compare the journal entries.",
        "Check doctor commission if a referral doctor was used.",
    ], "All amounts match. The bill and payment entries can be traced using their IDs.", "Invoice, receipts, ledger/journal and commission record.")
    scenario(doc, "T20", "Block cross-tenant and wrong-owner access", "Any two tenants, two doctors and two patients", "Prepare records in separate tenants and owners.", [
        "Use a record ID from Tenant B while signed in to Tenant A.",
        "Use Doctor A to request Doctor B's patient.",
        "Use Patient A to request Patient B's bill or report.",
    ], "Every request is denied. No name, report, bill or record-existence detail is leaked.", "Denied UI/API responses for all three checks.")

    doc.add_page_break()
    add_heading(doc, "9. Final end-to-end test", 1)
    add_body(doc, "Run this test after the individual scenarios pass.")
    numbered(doc, [
        "Admin creates and checks all required users.",
        "Front Desk registers a new patient.",
        "Front Desk creates a bill with an individual test and a package.",
        "Cashier records a partial payment.",
        "Phlebotomist collects and labels the samples.",
        "Lab Technician processes the samples and enters valid results.",
        "Pathologist reviews, approves and releases the report.",
        "Cashier records the final payment.",
        "Doctor opens the referred patient's Released report.",
        "Patient opens the bill, receipt and Released report.",
        "Accounts Manager checks the invoice, receipts, journal and commission.",
        "Inventory Manager checks stock and reservations.",
    ])
    add_callout(doc, "FINAL PASS", "All linked IDs are correct. The statuses are correct. The money totals match. Only authorized users can see the data.", "ok")
    add_heading(doc, "10. Testing completion checklist", 1)
    bullets(doc, [
        "All 20 scenarios have a result.",
        "Every failed scenario has a defect ID.",
        "Critical access, patient identity, report release and money defects are closed.",
        "Known product gaps are accepted or removed from the release scope.",
        "The final end-to-end test passes.",
    ])
    doc.save(TESTER_DOC)


def build_reference_doc(handoff_img):
    doc = Document()
    configure(doc, "LIMS | ROLES AND PRODUCT REFERENCE", "LIMS Roles, Gaps and Product Reference")
    title_page(
        doc,
        "PRODUCT REFERENCE",
        "LIMS Roles, Gaps and Product Reference",
        "Role boundaries, workflow rules, known risks and decisions",
        "Product owners, developers, business analysts and QA leads. This is not the tester step-by-step handbook.",
    )
    add_heading(doc, "1. Why this is a separate document", 1)
    add_body(doc, "The testing handbook contains only actions and expected results. This reference explains why the roles are separated, what the original flow missed, and which product decisions still need confirmation.")
    add_figure(doc, handoff_img, "Figure 1. Main clinical hand-off and report-release boundary.", 6.35)

    doc.add_page_break()
    add_heading(doc, "2. Correct role boundaries", 1)
    add_table(doc, ["Role", "Main responsibility", "Must not do by default"], [
        ("Platform/Developer Admin", "Create labs, plans, modules and tenant lifecycle", "Routine tenant clinical work"),
        ("Tenant Admin", "Users, roles, settings and full tenant oversight", "Bypass audit or workflow controls"),
        ("Lab Manager", "Manage daily lab work and release reports if policy allows", "Developer/platform work"),
        ("Front Desk Receptionist", "Patient registration and billing", "Clinical approval unless separately granted"),
        ("Billing Cashier", "Payments, dues, permitted refunds", "Result entry or report release"),
        ("Phlebotomist", "Collect, receive and reject samples", "Report approval or release"),
        ("Lab Technician", "Process samples and enter draft results", "Collect or release by default"),
        ("Report Typist", "Enter draft report values", "Verify or release"),
        ("Pathologist", "Review, approve and release reports", "Routine billing administration"),
        ("Accounts Manager", "Ledgers, expenses, closing and financial reports", "Clinical result work"),
        ("Inventory Manager", "Stock, suppliers, purchase and expiry", "Clinical approval"),
        ("Doctor Regular", "Owned referrals, Released reports and commission", "Other doctors' patients or draft reports"),
        ("Doctor Investor", "Approved investor views and order creation", "Unapproved tenant-wide data access"),
        ("Patient", "Own bills, receipts and Released reports", "Staff access or other patient data"),
    ], [2250, 3850, 3260], font_size=8.6)

    add_heading(doc, "3. Items missing from the original flow", 1)
    bullets(doc, [
        "Platform onboarding, subscription and enabled modules.",
        "Lab Manager, Phlebotomist, Report Typist, Pathologist, Cashier, Accounts and Inventory roles.",
        "Patient duplicate checks and repeat visits.",
        "Payment states, multiple payments, invoices, receipts, cancellation and refund.",
        "Barcode, custody history, rejection and recollection.",
        "Draft, Reviewed, Approved and Released report states.",
        "Doctor and patient activation, lockout and ownership rules.",
        "Accounts, commissions, stock movement and audit checks.",
        "Error paths and negative permission tests.",
    ])

    doc.add_page_break()
    add_heading(doc, "4. Status rules", 1)
    add_table(doc, ["Record", "Normal order", "Other state"], [
        ("Bill/order", "Open -> In progress -> Completed", "Cancelled"),
        ("Bill item", "Sample-pending -> Sample-collected -> Processing -> Reported", "Cancelled"),
        ("Sample", "Registered -> Collected -> Processing -> Completed -> Released", "Rejected or Archived"),
        ("Report", "Draft -> Reviewed -> Approved -> Released", "Correction process needs decision"),
        ("Payment", "Unpaid -> Partial -> Paid", "Cancelled/refunded rules need policy"),
    ], [1900, 4900, 2560], font_size=9.2)
    add_callout(doc, "Rule", "Clinical completion and payment completion are separate. The lab must decide whether unpaid or partially paid reports can be released.", "warn")

    add_heading(doc, "5. Current implementation risks", 1)
    bullets(doc, [
        "The default Lab Technician role cannot collect samples, verify reports or release reports.",
        "The current Front Desk template has broad permissions for test masters, samples and reports. Confirm whether this is intended.",
        "The report API checks sample names such as in-testing and testing-complete, but the Sample model uses Registered, Collected, Processing and Completed. This needs alignment.",
        "Doctor portal routing uses the linked doctor ID. The Doctor Regular template does not include the separate doctor-portal.access permission from the permission catalogue.",
        "Doctor Investor screens expose broad patient, billing and account information. Confirm privacy, plan and role restrictions.",
        "Automatic inventory consumption, barcode scanning/printing, rejection approval, full refund/daily close, complete audit coverage and post-release correction are not fully complete.",
    ])

    add_heading(doc, "6. Product decisions required", 1)
    add_table(doc, ["Decision", "Question to answer"], [
        ("Front Desk permissions", "Should reception edit reports, collect samples or create/delete test masters?"),
        ("Sample collection", "Can a Lab Technician collect, or only a Phlebotomist/authorized Front Desk user?"),
        ("Report release", "Must a Pathologist release, or can Lab Manager/Admin also release?"),
        ("Payment gate", "Can an unpaid or partially paid report be released?"),
        ("Investor access", "Can an Investor view all patients, bills and accounts and create patient/orders?"),
        ("Recollection", "Does rejection create a new sample while keeping the old sample as Rejected?"),
        ("Report correction", "How is a Released report corrected, approved and shown in portals?"),
        ("Stock enforcement", "Should low, insufficient or expired stock warn or block testing?"),
    ], [2800, 6560], font_size=9.1)

    add_heading(doc, "7. Recommended product rule", 1)
    add_callout(doc, "Recommended", "Keep the roles separated. Use custom roles only when a laboratory truly needs one person to perform more than one job. Always enforce the same permission in both the screen and the API.", "ok")
    add_body(doc, "When a decision changes, update the role template, server permission, screen visibility, workflow diagram and test scenario together.")
    doc.save(REFERENCE_DOC)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)
    flow = ASSETS / "simple_testing_flow.png"
    handoff = ASSETS / "simple_role_handoff.png"
    save_simple_flow(flow)
    save_handoff_flow(handoff)
    build_tester_doc(flow, handoff)
    build_reference_doc(handoff)
    print(TESTER_DOC)
    print(REFERENCE_DOC)


if __name__ == "__main__":
    main()
