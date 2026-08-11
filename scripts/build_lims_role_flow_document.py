from pathlib import Path
from textwrap import wrap
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "deliverables"
ASSET_DIR = OUT_DIR / "lims_role_flow_assets"
DOCX_PATH = OUT_DIR / "LIMS_Role_Workflow_and_Testing_Guide.docx"

NAVY = "163A5F"
BLUE = "2E74B5"
LIGHT_BLUE = "E8F1F8"
PALE = "F4F6F9"
GRAY = "64748B"
LIGHT_GRAY = "E5E7EB"
GREEN = "DCFCE7"
GREEN_DARK = "166534"
AMBER = "FEF3C7"
AMBER_DARK = "92400E"
RED = "FEE2E2"
RED_DARK = "991B1B"
WHITE = "FFFFFF"
BLACK = "111827"


def font(size=28, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def hex_rgb(value):
    value = value.lstrip("#")
    return tuple(int(value[i:i+2], 16) for i in (0, 2, 4))


def rounded_box(draw, box, fill, outline=BLUE, radius=18, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=hex_rgb(fill), outline=hex_rgb(outline), width=width)


def draw_centered_text(draw, box, text, size=25, bold=False, color=BLACK, max_chars=28, line_gap=7):
    f = font(size, bold)
    lines = []
    for para in text.split("\n"):
        lines.extend(wrap(para, max_chars) or [""])
    heights = []
    widths = []
    for line in lines:
        b = draw.textbbox((0, 0), line, font=f)
        widths.append(b[2] - b[0])
        heights.append(b[3] - b[1])
    total_h = sum(heights) + max(0, len(lines)-1) * line_gap
    y = box[1] + (box[3] - box[1] - total_h) / 2
    for i, line in enumerate(lines):
        x = box[0] + (box[2] - box[0] - widths[i]) / 2
        draw.text((x, y), line, font=f, fill=hex_rgb(color))
        y += heights[i] + line_gap


def arrow(draw, start, end, color=GRAY, width=5, head=15):
    draw.line([start, end], fill=hex_rgb(color), width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2-x1) >= abs(y2-y1):
        s = 1 if x2 > x1 else -1
        pts = [(x2, y2), (x2-s*head, y2-head//2), (x2-s*head, y2+head//2)]
    else:
        s = 1 if y2 > y1 else -1
        pts = [(x2, y2), (x2-head//2, y2-s*head), (x2+head//2, y2-s*head)]
    draw.polygon(pts, fill=hex_rgb(color))


def save_end_to_end_diagram(path):
    w, h = 1800, 2300
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    d.text((80, 45), "LIMS end-to-end operational flow", font=font(42, True), fill=hex_rgb(NAVY))
    d.text((80, 105), "Owner shown in each step; exception paths are highlighted.", font=font(25), fill=hex_rgb(GRAY))

    steps = [
        ("ADMIN / LAB MANAGER", "Configure lab, modules, users, roles and master data", LIGHT_BLUE),
        ("FRONT DESK", "Find or register patient; select referral doctor", "EEF2FF"),
        ("FRONT DESK / CASHIER", "Create bill/order, select tests, priority, discounts and payment", "EEF2FF"),
        ("SYSTEM", "Create linked sample record(s), IDs and barcodes", PALE),
        ("PHLEBOTOMIST", "Verify patient + order, collect/receive and label specimen", "ECFDF5"),
        ("LAB TECHNICIAN", "Process specimen, reserve/monitor inventory, enter results", "ECFDF5"),
        ("PATHOLOGIST / LAB MANAGER", "Review -> approve -> release report", AMBER),
        ("SYSTEM / FRONT DESK", "Notify, print/download, settle dues and post accounts", PALE),
        ("DOCTOR / PATIENT", "View only authorized, released reports and permitted billing data", LIGHT_BLUE),
    ]
    x1, x2 = 250, 1550
    top, box_h, gap = 180, 165, 64
    prev = None
    for idx, (role, text, fill) in enumerate(steps):
        y1 = top + idx*(box_h+gap)
        box = (x1, y1, x2, y1+box_h)
        rounded_box(d, box, fill)
        d.rounded_rectangle((x1+18, y1+22, x1+410, y1+box_h-22), radius=14, fill=hex_rgb(NAVY))
        draw_centered_text(d, (x1+28, y1+28, x1+400, y1+box_h-28), role, size=23, bold=True, color=WHITE, max_chars=24)
        draw_centered_text(d, (x1+440, y1+18, x2-24, y1+box_h-18), text, size=26, bold=False, max_chars=42)
        if prev:
            arrow(d, prev, ((x1+x2)//2, y1-4), color=BLUE)
        prev = ((x1+x2)//2, y1+box_h+4)

    # Exception branches.
    y_sample = top + 4*(box_h+gap) + box_h//2
    rounded_box(d, (20, y_sample-65, 215, y_sample+65), RED, RED_DARK, radius=14)
    draw_centered_text(d, (25, y_sample-60, 210, y_sample+60), "Reject / recollect\nwith reason", 20, True, RED_DARK, 18)
    arrow(d, (250, y_sample), (215, y_sample), RED_DARK, width=4)
    y_report = top + 6*(box_h+gap) + box_h//2
    rounded_box(d, (1585, y_report-65, 1780, y_report+65), AMBER, AMBER_DARK, radius=14)
    draw_centered_text(d, (1590, y_report-60, 1775, y_report+60), "Return draft\nfor correction", 20, True, AMBER_DARK, 18)
    arrow(d, (1550, y_report), (1585, y_report), AMBER_DARK, width=4)
    im.save(path, quality=95)


def save_role_map(path):
    w, h = 1800, 1600
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 45), "Role interaction and access boundaries", font=font(42, True), fill=hex_rgb(NAVY))
    center = (620, 610, 1180, 920)
    rounded_box(d, center, LIGHT_BLUE, NAVY, radius=28, width=5)
    draw_centered_text(d, center, "LIMS CORE\nPatient -> Billing -> Sample -> Report", 31, True, NAVY, 30)
    roles = [
        ((80, 230, 540, 450), "ADMIN / LAB MANAGER", "Tenant setup, RBAC, masters, oversight"),
        ((1260, 230, 1720, 450), "FRONT DESK / CASHIER", "Patient, order, invoice, receipt and dues"),
        ((80, 1080, 540, 1330), "PHLEBOTOMIST / TECHNICIAN", "Collect, process and enter draft results"),
        ((1260, 1080, 1720, 1330), "PATHOLOGIST", "Verify, approve, release and print"),
        ((70, 610, 500, 850), "DOCTOR", "Owned referrals; released reports; commissions"),
        ((1300, 610, 1730, 850), "PATIENT", "Own bills/receipts and released reports only"),
    ]
    for box, role, desc in roles:
        rounded_box(d, box, PALE, BLUE, radius=20)
        draw_centered_text(d, (box[0]+15, box[1]+15, box[2]-15, box[1]+90), role, 23, True, NAVY, 27)
        draw_centered_text(d, (box[0]+22, box[1]+90, box[2]-22, box[3]-15), desc, 23, False, BLACK, 30)
        cx = (box[0]+box[2])//2
        cy = (box[1]+box[3])//2
        tx = center[0] if cx < center[0] else center[2]
        ty = min(max(cy, center[1]+35), center[3]-35)
        sx = box[2] if cx < center[0] else box[0]
        arrow(d, (sx, cy), (tx, ty), BLUE, width=4)
    d.rounded_rectangle((220, 1450, 1580, 1535), radius=15, fill=hex_rgb(RED), outline=hex_rgb(RED_DARK), width=3)
    draw_centered_text(d, (230, 1455, 1570, 1530), "Boundary rule: draft/reviewed/approved reports stay internal; portals receive released reports only.", 24, True, RED_DARK, 100)
    im.save(path, quality=95)


def save_status_lifecycle(path):
    w, h = 1800, 1500
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 45), "Operational status lifecycles", font=font(42, True), fill=hex_rgb(NAVY))

    lanes = [
        ("BILL / ORDER", ["Open", "In progress", "Completed"], 230, LIGHT_BLUE),
        ("SAMPLE", ["Registered", "Collected", "Processing", "Completed", "Released"], 600, GREEN),
        ("REPORT", ["Draft", "Reviewed", "Approved", "Released"], 970, AMBER),
        ("PAYMENT", ["Unpaid", "Partial", "Paid"], 1320, "EEF2FF"),
    ]
    for label, states, y, fill in lanes:
        d.text((70, y+34), label, font=font(25, True), fill=hex_rgb(NAVY))
        start_x, end_x = 300, 1690
        gap = 35
        bw = int((end_x-start_x-gap*(len(states)-1))/len(states))
        boxes = []
        for i, state in enumerate(states):
            x1 = start_x + i*(bw+gap)
            box = (x1, y, x1+bw, y+120)
            boxes.append(box)
            rounded_box(d, box, fill, BLUE, radius=16)
            draw_centered_text(d, box, state, 25, True, NAVY, 18)
            if i:
                arrow(d, (boxes[i-1][2]+3, y+60), (x1-3, y+60), BLUE, width=4, head=12)
        if label in ("BILL / ORDER", "PAYMENT"):
            d.text((start_x, y+145), "Alternate terminal state: Cancelled", font=font(22), fill=hex_rgb(RED_DARK))
        if label == "SAMPLE":
            d.text((start_x, y+145), "Exception from active states: Rejected (reason required); archival is separate.", font=font(22), fill=hex_rgb(RED_DARK))
    im.save(path, quality=95)


def save_portal_flows(path):
    w, h = 1800, 1800
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 45), "Doctor and patient portal flows", font=font(42, True), fill=hex_rgb(NAVY))

    def lane(x1, x2, title, steps, fill):
        d.rounded_rectangle((x1, 130, x2, 1720), radius=24, fill=hex_rgb("F8FAFC"), outline=hex_rgb(LIGHT_GRAY), width=3)
        d.text((x1+35, 170), title, font=font(32, True), fill=hex_rgb(NAVY))
        top, bh, gap = 275, 190, 65
        prev = None
        for idx, text in enumerate(steps):
            y1 = top + idx*(bh+gap)
            box = (x1+55, y1, x2-55, y1+bh)
            rounded_box(d, box, fill, BLUE, radius=18)
            draw_centered_text(d, box, f"{idx+1}. {text}", 24, idx == 0, BLACK, 34)
            if prev:
                arrow(d, prev, ((x1+x2)//2, y1-5), BLUE, width=4)
            prev = ((x1+x2)//2, y1+bh+5)

    lane(60, 870, "DOCTOR", [
        "Staff creates doctor + linked invited user",
        "Doctor activates by email OTP and creates password",
        "Doctor signs in; ownership is resolved from the session",
        "View referred patients, released reports and commission data",
        "Optionally create patient/test request; lab fulfils it",
    ], LIGHT_BLUE)
    lane(930, 1740, "PATIENT", [
        "Staff registers patient and prints private QR/PIN slip",
        "Patient scans QR, confirms access PIN + DOB, sets portal PIN",
        "Patient signs in with patient ID, DOB and portal PIN",
        "View own bills, receipts and released reports only",
        "Reissue slip to invalidate prior credentials and sessions",
    ], GREEN)
    im.save(path, quality=95)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run(run, size=11, bold=False, color=BLACK, italic=False, font_name="Calibri"):
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    keep_with_next(p)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run(r)
    else:
        r = p.add_run(text)
        set_run(r)
    return p


def add_bullets(doc, items, style="List Bullet"):
    for item in items:
        p = doc.add_paragraph(style=style)
        r = p.add_run(item)
        set_run(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run(r)


def add_callout(doc, label, text, kind="info"):
    palette = {
        "info": (LIGHT_BLUE, NAVY),
        "warn": (AMBER, AMBER_DARK),
        "risk": (RED, RED_DARK),
        "ok": (GREEN, GREEN_DARK),
    }
    fill, color = palette[kind]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.10)
    p.paragraph_format.right_indent = Inches(0.10)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(f"{label}: ")
    set_run(r, bold=True, color=color)
    r = p.add_run(text)
    set_run(r, color=color)


def add_table(doc, headers, rows, widths, header_fill="E8EEF5", font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run(r, size=font_size)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc, image_path, caption, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption.split(".", 1)[0])
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_before = Pt(2)
    c.paragraph_format.space_after = Pt(8)
    r = c.add_run(caption)
    set_run(r, size=9, italic=True, color=GRAY)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, before, after, color in [
        (1, 16, 18, 10, BLUE), (2, 13, 14, 7, BLUE), (3, 12, 10, 5, "1F4D78")
    ]:
        st = styles[f"Heading {level}"]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.188)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("LIMS ROLE WORKFLOW | TESTING HANDOFF")
    set_run(r, size=8.5, bold=True, color=GRAY)
    add_page_number(section.footer.paragraphs[0])

    doc.core_properties.title = "LIMS Role Workflow and Testing Guide"
    doc.core_properties.subject = "End-to-end role flows, access boundaries and QA guidance"
    doc.core_properties.author = "LIMS Product Team"
    doc.core_properties.keywords = "LIMS, workflow, roles, testing, QA, UAT"


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "end": ASSET_DIR / "end_to_end.png",
        "roles": ASSET_DIR / "role_map.png",
        "status": ASSET_DIR / "status_lifecycle.png",
        "portals": ASSET_DIR / "portal_flows.png",
    }
    save_end_to_end_diagram(diagrams["end"])
    save_role_map(diagrams["roles"])
    save_status_lifecycle(diagrams["status"])
    save_portal_flows(diagrams["portals"])

    doc = Document()
    configure_document(doc)

    # Cover / customer-pack opening.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("QUALITY ASSURANCE HANDOFF")
    set_run(r, size=10, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("LIMS Role Workflow and Testing Guide")
    set_run(r, size=26, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Corrected end-to-end process, role boundaries, status lifecycles, portal access and test focus areas")
    set_run(r, size=13, color=GRAY)

    add_table(doc, ["Document purpose", "Basis", "Prepared for"], [[
        "Manual QA, UAT and automation planning",
        "Provided flow image + repository implementation reviewed on 11 Aug 2026",
        "LIMS testing team",
    ]], [3000, 3900, 2460], font_size=9.5)
    add_callout(doc, "Key correction", "Sample collection, result entry and report release are separate responsibilities in the current RBAC model. Phlebotomist collects; Lab Technician/Report Typist enters results; Pathologist or Lab Manager verifies and releases.", "warn")
    add_figure(doc, diagrams["roles"], "Figure 1. Role interaction and access boundaries.", 6.25)

    doc.add_page_break()
    add_heading(doc, "1. Executive analysis: what was missing or unclear", 1)
    add_body(doc, "The supplied flow captures the central happy path: administration, patient registration, test ordering, billing, specimen work, reporting, doctor access and patient access. The following additions are necessary for a complete and testable workflow.")
    missing_rows = [
        ("Role coverage", "Add Platform/Developer Admin, Lab Manager, Phlebotomist, Report Typist, Pathologist, Billing Cashier, Accounts Manager and Inventory Manager. Patient portal users are not tenant staff roles."),
        ("Role separation", "Do not assign collection + result entry + verification + release to one generic Lab Technician unless a custom role explicitly grants all permissions."),
        ("Doctor type vs access role", "Keep profile/business type (Investor or Regular) separate from authentication role and referral ownership. Confirm the intended Investor scope."),
        ("Preconditions", "Add lab onboarding, enabled modules/subscription, role seeding, users, branding/settings, test master and inventory master before transactions."),
        ("Patient intake", "Add patient lookup/duplicate handling, demographics, referring doctor, identity confirmation, repeat visit and portal-access issuance."),
        ("Billing", "Add priority, package expansion, discounts/tax permissions, unpaid/partial/paid states, multiple payments, invoice/receipt, cancellation/refund and corporate credit."),
        ("Sample lifecycle", "Add automatic sample creation from billing, barcode/label, patient/specimen verification, received time, custody log, rejection/recollection and status synchronization."),
        ("Inventory", "Test reservation, stock movement, low/reorder/expiry alerts and insufficient-stock behavior; automatic consumption remains an implementation gap."),
        ("Report lifecycle", "Add draft -> reviewed -> approved -> released, abnormal flags, correction return, version history, print permission and released-only portal visibility."),
        ("Doctor portal", "Add invitation/activation, server-side referral ownership, released-only reports, commission eligibility and payout history."),
        ("Patient portal", "Add confidential QR/PIN issuance, DOB validation, portal PIN, lockout/session invalidation, receipts and released-only reports."),
        ("Cross-cutting controls", "Add tenant isolation, negative RBAC tests, audit/notification checks, session lockout, data consistency, cancellation effects and error/retry behavior."),
    ]
    add_table(doc, ["Area", "Required clarification/addition"], missing_rows, [2150, 7210], font_size=9.1)

    add_heading(doc, "1.1 Implementation risks the testing team should confirm", 2)
    add_bullets(doc, [
        "The Lab Technician template can update samples and edit reports, but it cannot collect samples, verify reports or release reports.",
        "The Front Desk Receptionist template currently includes test create/delete, sample collect/update and report edit/print permissions, which is broader than its description and should be confirmed as intentional.",
        "The report-creation API checks sample states named in-testing and testing-complete, while the Sample model uses registered, collected, processing, completed, released, rejected and archived. Exercise this integration because the vocabulary is inconsistent.",
        "Doctor Regular portal routing is driven by the linked doctorId and ownership checks; the role template does not include the explicit doctor-portal.access catalog permission. Confirm the intended permission contract.",
        "The Doctor Investor UI exposes broad lab analytics, billing, accounts and all-patient views. Treat this as privileged access and validate that it is limited to the intended investor role and plan.",
        "Automatic inventory consumption, barcode scanning/printing polish, rejection approval, full refund/daily-close controls, comprehensive audit coverage and report correction approval are documented as incomplete areas.",
    ])
    add_callout(doc, "Testing interpretation", "Items above are not assumptions that the feature is complete. They are explicit verification targets and product decisions that need sign-off.", "risk")

    doc.add_page_break()
    add_heading(doc, "2. Corrected end-to-end flow", 1)
    add_figure(doc, diagrams["end"], "Figure 2. Corrected operational workflow with exception paths.", 6.25)

    add_heading(doc, "2.1 Happy path", 2)
    add_numbered(doc, [
        "Platform/Developer Admin onboards the lab tenant, activates the subscription/modules and provisions the initial tenant Admin.",
        "Tenant Admin or Lab Manager creates/assigns staff roles, configures lab settings, doctors, test categories/definitions/packages, parameters/reference ranges and inventory masters.",
        "Front Desk searches for an existing patient; if none is found, registers the patient and captures the referring doctor where applicable.",
        "Front Desk or Billing Cashier selects tests/packages, priority, permitted discount/tax/payment details and creates the bill/order.",
        "The system assigns the bill ID, creates linked sample records with sample IDs/barcodes, posts the invoice/receivable entry and exposes the order in sample worklists.",
        "Phlebotomist confirms two patient identifiers and specimen requirements, collects/receives and labels the specimen, then advances it to processing or rejects it with a reason.",
        "Lab Technician processes the specimen, monitors/reserves required inventory, records parameter results and completes the sample; a draft report is created/updated.",
        "Report Typist/Lab Technician completes draft values. Pathologist or Lab Manager reviews, approves and releases the report in sequence.",
        "On release, authorized staff can print/download; doctor and patient portals can retrieve only their authorized released report. Billing, receipts, commissions, notifications and accounts reflect the transaction.",
    ])

    add_heading(doc, "2.2 Exception paths", 2)
    add_bullets(doc, [
        "Duplicate/incorrect patient: stop order creation, select the correct existing record or correct data with permission; preserve linked-record integrity.",
        "Inactive doctor/test/package or invalid price: block selection and avoid partial records.",
        "Insufficient/expired inventory: warn or block according to configured policy; do not silently consume negative stock.",
        "Specimen mismatch, insufficient volume, hemolysis or invalid timing: reject with reason, retain custody history and create a controlled recollection path.",
        "Invalid/missing result: prevent review; return draft for correction where the review fails.",
        "Cancelled bill/refund: verify sample/report restrictions, payment reversal, receipt/journal behavior and audit entries.",
        "Portal ownership failure or non-released report: return no medical/report data; avoid record-existence disclosure.",
    ])

    doc.add_page_break()
    add_heading(doc, "3. Role catalogue and expected boundaries", 1)
    role_rows = [
        ("Platform/Developer Admin", "Master platform", "Onboard labs; plans/modules; tenant lifecycle; role templates", "No routine tenant clinical processing"),
        ("Tenant Admin", "Tenant", "Full tenant setup, users, roles, settings and oversight", "Dangerous actions require strong audit"),
        ("Lab Manager", "Tenant", "Operational management; samples; verify/release; billing collect; reports/accounts view", "No developer/platform access"),
        ("Front Desk Receptionist", "Tenant", "Patient/doctor lookup; registration; billing; permitted sample/report actions", "Confirm currently broad template permissions"),
        ("Billing Cashier", "Tenant", "Bills, payments, dues, discounts/refunds and accounts view", "No result approval/release"),
        ("Phlebotomist", "Tenant", "Sample collect/update/reject; patient and order lookup", "No result approval/release"),
        ("Lab Technician", "Tenant", "Process samples; tests; draft report result entry", "No collection, verification or release by default"),
        ("Report Typist", "Tenant", "Enter draft report values", "Cannot verify or release"),
        ("Pathologist", "Tenant", "Review, approve, release and print reports", "No routine billing/refund administration"),
        ("Accounts Manager", "Tenant", "Collections/refunds; ledgers; expenses; reports; closing", "No clinical result entry/release"),
        ("Inventory Manager", "Tenant", "Items, suppliers, UOMs, stock movements, purchase/expiry", "No clinical/report approval"),
        ("Doctor Regular", "Portal", "Owned referrals, visits, released reports, commission/profile", "No other doctors' patients; no draft reports"),
        ("Doctor Investor", "Privileged portal", "Investor analytics and configured broader lab views/order creation", "Scope must be expressly approved and plan-gated"),
        ("Patient", "Separate portal", "Own bills/receipts and released reports", "No staff RBAC; no other patient or draft data"),
    ]
    add_table(doc, ["Role", "Scope", "Expected capabilities", "Key boundary"], role_rows, [1900, 1050, 3850, 2560], font_size=8.2)

    add_heading(doc, "3.1 Recommended segregation of duties", 2)
    add_table(doc, ["Activity", "Performer", "Independent control"], [
        ("Create patient and order", "Front Desk / authorized Doctor", "Duplicate and identity validation"),
        ("Collect specimen", "Phlebotomist", "Patient/order/label match"),
        ("Enter result", "Lab Technician / Report Typist", "Required values and range validation"),
        ("Verify/approve/release", "Pathologist / Lab Manager", "Status sequence and release permission"),
        ("Refund/cancel/close", "Cashier / Accounts Manager", "Reason, limits and audit/reconciliation"),
        ("Role/permission change", "Tenant Admin", "Dangerous-grant review and audit"),
    ], [3000, 3200, 3160], font_size=9.1)

    doc.add_page_break()
    add_heading(doc, "4. Detailed role flows", 1)
    flows = [
        ("4.1 Platform/Developer Admin", [
            "Authenticate in the developer portal.", "Create/activate the lab tenant and subscription; configure enabled modules.", "Provision or reset the initial tenant Admin.", "Maintain platform permission catalog/role templates and tenant lifecycle.", "Verify tenant isolation, plan/quota enforcement and auditability."
        ]),
        ("4.2 Tenant Admin / Lab Manager", [
            "Sign in to the correct tenant.", "Create roles/users or apply seeded templates; set active/inactive state.", "Configure lab profile, test master, doctors and inventory masters.", "Monitor dashboard, samples, reports, billing, accounts and inventory within permission/plan scope.", "Review audit and exceptions; do not bypass clinical/financial status sequences."
        ]),
        ("4.3 Front Desk / Billing Cashier", [
            "Search patient first; register only if no safe match exists.", "Capture/update demographics, referral doctor and visit details.", "Select active tests/packages and priority; calculate totals, discount/tax and commission snapshot.", "Create bill/order, print invoice and collect full/partial/no payment; issue receipt.", "Explain sample next step and optionally issue patient portal access after identity confirmation.", "For cancellation/refund, require permission and reason; verify downstream and accounting effects."
        ]),
        ("4.4 Phlebotomist", [
            "Open pending sample worklist and confirm patient + bill + required specimen.", "Use two patient identifiers; verify container, fasting/timing and label/barcode.", "Record collection/received time and collector/receiver; mark Collected then Processing.", "If unacceptable, reject with reason and initiate recollection without erasing custody history.", "Confirm bill-item/sample statuses stay synchronized."
        ]),
        ("4.5 Lab Technician / Report Typist", [
            "Open a collected/processing specimen assigned for testing.", "Verify equipment/QC and required inventory availability; reserve/record consumption where supported.", "Enter every required numeric/text result; verify flags, units and ranges.", "Complete the sample and save a Draft report.", "Respond to a returned draft; do not verify or release without the relevant permission."
        ]),
        ("4.6 Pathologist / Lab Manager", [
            "Open Draft report and verify patient, sample, test, results, units, ranges, flags and remarks.", "Review Draft -> Reviewed; approve Reviewed -> Approved; release Approved -> Released.", "Return incorrect work before approval/release and preserve version/audit data.", "Confirm released output is printable and becomes visible to the correct doctor/patient only.", "Treat post-release correction/amendment as a controlled gap requiring product sign-off."
        ]),
        ("4.7 Inventory Manager", [
            "Maintain categories/types/UOMs/suppliers/locations/storage conditions and items.", "Load stock through controlled movements or purchase orders; capture batch and expiry.", "Monitor on-hand, reserved, minimum/reorder and expiring/expired quantities.", "Confirm test-to-inventory requirements and insufficient-stock behavior.", "Reconcile physical and system stock; investigate negative/unexpected movement."
        ]),
    ]
    for title, steps in flows:
        add_heading(doc, title, 2)
        add_numbered(doc, steps)

    doc.add_page_break()
    add_heading(doc, "5. Doctor and patient flows", 1)
    add_figure(doc, diagrams["portals"], "Figure 3. Doctor and patient portal activation and use.", 6.25)
    add_heading(doc, "5.1 Doctor Regular", 2)
    add_bullets(doc, [
        "Staff creates the doctor profile with a unique email; the system creates a linked invited user and sends a six-digit activation OTP.",
        "The doctor activates within the invitation window, creates a password and later signs in on the correct tenant.",
        "Server-side ownership is derived from the signed session/doctorId. Only referred patients and released reports are returned.",
        "Commission is estimated until eligible; paid referral rules determine earned commission. Payouts remain staff-controlled.",
        "Inactive doctor, expired/used OTP, wrong tenant, non-owned patient and draft report requests must be denied.",
    ])
    add_heading(doc, "5.2 Doctor Investor", 2)
    add_bullets(doc, [
        "Validate that the user is explicitly assigned the Investor role and eligible subscription plan.",
        "Confirm access to investor analytics, lab billings/accounts/all-patient views and the ability to register patients/send test requests exactly matches the approved policy.",
        "A doctor-created request must enter the lab work queue; staff still controls billing settlement, collection, processing and report release.",
        "Test that a Regular doctor cannot reach Investor APIs/pages by URL or crafted request.",
    ])
    add_heading(doc, "5.3 Patient", 2)
    add_bullets(doc, [
        "Staff confirms identity before printing/replacing the confidential QR/PIN access slip.",
        "Activation requires QR token, six-digit access PIN and DOB, then a private four-digit portal PIN is set.",
        "Returning login uses Patient ID + DOB + portal PIN; five failed attempts temporarily lock access.",
        "The patient sees only their own released reports, bill totals/status, receipts and safe result parameters.",
        "Reissuing a slip invalidates older activation data and prior patient sessions.",
    ])

    doc.add_page_break()
    add_heading(doc, "6. Status model and synchronization", 1)
    add_figure(doc, diagrams["status"], "Figure 4. Separate but linked order, sample, report and payment lifecycles.", 6.25)
    add_callout(doc, "Important", "Clinical completion and financial settlement are independent. A report may progress while a bill remains unpaid/partial unless the laboratory adopts an explicit payment-before-release policy.", "warn")
    add_table(doc, ["Event", "Expected synchronized effect"], [
        ("Bill created", "Bill Open; items Sample-pending; sample record(s) Registered; invoice/receivable posted."),
        ("Sample collected", "Sample Collected; linked bill item Sample-collected; custody timestamp/user retained."),
        ("Processing starts", "Sample Processing; bill/order In progress; inventory reservation/alerts evaluated."),
        ("Results completed", "Sample Completed; bill items become Reported when report is generated; report remains Draft."),
        ("Report released", "Report Released; sample may become Released; portals can access subject to ownership."),
        ("Payment collected", "Billing Unpaid -> Partial/Paid; receipt and journal/payment history created."),
        ("Cancellation/refund", "Financial reversal/receipt history and downstream restrictions follow policy; audit reason required."),
    ], [2500, 6860], font_size=9.1)

    doc.add_page_break()
    add_heading(doc, "7. Tester-ready scenario checklist", 1)
    scenario_rows = [
        ("SET-01", "Seed/create every role and assign a user", "Correct navigation and API permissions"),
        ("SET-02", "Disable user/doctor and retry active session/login", "Access denied promptly"),
        ("RBAC-01", "Technician attempts collection/release", "Denied without permission"),
        ("RBAC-02", "Regular doctor requests Investor page/API", "Denied; no data leakage"),
        ("PAT-01", "Search then register a new patient", "Unique ID; portal account behavior correct"),
        ("PAT-02", "Attempt likely duplicate patient", "Warning/block per policy; no duplicate links"),
        ("BILL-01", "Create order with tests/package/referral/priority", "Totals, snapshots, samples, accounts correct"),
        ("BILL-02", "Partial then final multi-mode payment", "Due/status/receipts/journals correct"),
        ("BILL-03", "Overpayment/unauthorized discount/refund", "Blocked and audited"),
        ("SMP-01", "Collect -> process -> complete", "Valid sequence, timestamps, custody and bill sync"),
        ("SMP-02", "Reject from each active sample state", "Reason required; terminal rejection; recollection handled"),
        ("SMP-03", "Barcode/label mismatch", "Wrong patient/specimen cannot proceed"),
        ("INV-01", "Reserve/consume required items", "Stock/reserved/UOM conversion accurate"),
        ("INV-02", "Low/expired/insufficient stock", "Configured warning/block; no negative silent stock"),
        ("REP-01", "Enter missing/invalid/exponential result", "Validation prevents unsafe draft/review"),
        ("REP-02", "Draft -> Reviewed -> Approved -> Released", "Only valid sequential transitions"),
        ("REP-03", "Portal reads non-released report", "404/denial and zero medical leakage"),
        ("DOC-01", "Invite, activate, resend, expire doctor OTP", "Single-use behavior and old OTP invalidation"),
        ("DOC-02", "Doctor opens non-owned referral", "404/denial; ownership enforced server-side"),
        ("DOC-03", "Commission estimate/earn/payout", "Historical rate snapshot and accounting reference"),
        ("PORT-01", "Patient QR/PIN activation and returning login", "Correct credential/session lifecycle"),
        ("PORT-02", "Five failures then reissue slip", "Lockout and old-session invalidation"),
        ("TEN-01", "Cross-tenant ID/API probing", "No existence disclosure or data crossover"),
        ("AUD-01", "Sensitive actions across modules", "Actor/action/time/object/reason captured where implemented"),
    ]
    add_table(doc, ["ID", "Scenario", "Expected result"], scenario_rows, [1150, 4900, 3310], font_size=8.1)

    add_heading(doc, "7.1 End-to-end UAT data set", 2)
    add_bullets(doc, [
        "One new and one returning patient; one duplicate candidate; one minor/guardian case if supported.",
        "One Regular and one Investor doctor; active/inactive variants; different commission rates.",
        "Individual tests and package with numeric/text parameters, required/optional values, normal and abnormal results.",
        "Routine and urgent orders; unpaid, partial and paid bills; cash/UPI/card/corporate variants.",
        "Valid, low-stock and expired inventory batches with UOM conversion.",
        "Samples that complete normally, reject, require recollection and contain a label mismatch.",
    ])

    doc.add_page_break()
    add_heading(doc, "8. Test evidence and exit criteria", 1)
    add_heading(doc, "8.1 Evidence to retain", 2)
    add_bullets(doc, [
        "User/role/tenant/plan used; test-data IDs; date/time and environment/build identifier.",
        "Before/after screenshots for each state change; API response for access-control negatives.",
        "Bill, sample, report, receipt/journal and inventory identifiers demonstrating linkage.",
        "Audit/notification entries where implemented; note explicitly when expected coverage is absent.",
        "Portal proof showing released-only and ownership filters without exposing unrelated patient data.",
    ])
    add_heading(doc, "8.2 Exit criteria", 2)
    add_bullets(doc, [
        "Every role can complete its authorized happy path and cannot complete unauthorized actions by UI or API.",
        "Patient -> bill -> sample -> report links remain correct through normal, rejected, cancelled and partially paid paths.",
        "Only released reports reach doctor/patient portals, and ownership/tenant boundaries hold under direct URL/API probes.",
        "Payment totals, receipts, commissions, journals and inventory movements reconcile for the UAT data set.",
        "No unresolved Critical/High defects in authorization, tenant isolation, patient identity, report release or money movement.",
        "Known incomplete features are documented with an accepted workaround, release exclusion or implementation commitment.",
    ])

    add_heading(doc, "9. Decisions required before final QA sign-off", 1)
    decision_rows = [
        ("Front Desk privileges", "Should reception create/delete test masters, edit reports and collect/update samples?", "Approve narrower template or retain with rationale."),
        ("Technician collection", "Should Lab Technician collect specimens, or only Phlebotomist/Front Desk?", "Define standard lab staffing model and custom-role guidance."),
        ("Report release", "Is Pathologist mandatory, or may Lab Manager/Admin release?", "Define approval segregation and emergency override."),
        ("Payment gate", "Can unpaid/partial bills release reports?", "Define lab policy and enforcement point."),
        ("Investor scope", "May Investor see all patients, bills/accounts and create patient/orders?", "Document consent, plan and privacy boundary."),
        ("Recollection", "Does rejection create a new sample or reset the old one?", "Preserve chain of custody and billing linkage."),
        ("Post-release correction", "How are amended reports approved and communicated?", "Define version, reason, approver and portal behavior."),
        ("Inventory enforcement", "Warn or block when stock is insufficient/expired?", "Define override authority and audit."),
    ]
    add_table(doc, ["Decision", "Question", "Required outcome"], decision_rows, [1900, 4650, 2810], font_size=8.7)

    add_heading(doc, "Appendix A. Source and interpretation notes", 1)
    add_body(doc, "This guide was derived from the user-provided role-flow image and validated against the repository's current RBAC catalog, sidebar/portal routing, billing/sample/report models and APIs, doctor and patient portal notes, module documentation, and existing test scenario files. Where repository documentation labels a capability as pending or where code contracts conflict, the guide identifies it as a testing risk or decision rather than claiming completion.")
    add_body(doc, "Status names and permissions should be treated as the implementation baseline for this build. If product policy changes, update the role templates, APIs, diagrams and tests together to prevent UI/API drift.")

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
