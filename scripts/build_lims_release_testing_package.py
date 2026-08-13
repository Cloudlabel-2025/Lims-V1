from pathlib import Path
import json
from collections import defaultdict

from docx import Document
from docx.shared import Pt

from build_lims_split_documents import configure, title_page, numbered, bullets, label_para
from build_lims_role_flow_document import add_heading, add_body, add_callout, add_table


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "deliverables"
QA = ROOT / "docs" / ".qa_release_package"
CATALOG = QA / "lims_release_test_cases.json"
DOCS = {
    1: OUT / "LIMS_Level_1_Core_Flow_Testing.docx",
    2: OUT / "LIMS_Level_2_Functional_Integration_Testing.docx",
    3: OUT / "LIMS_Level_3_Market_Readiness_Testing.docx",
}


CASES = []
COUNTERS = defaultdict(int)


def add_case(level, module, title, action, expected, role, test_type="Functional",
             priority="Medium", gap="", gate=False, precondition="Required master data and user exist.",
             test_data="Use dedicated non-production test data."):
    COUNTERS[level] += 1
    cid = f"L{level}-{COUNTERS[level]:03d}"
    steps = f"1. Sign in as {role}. 2. {action} 3. Save or submit. 4. Refresh or reopen the record. 5. Compare the result with the pass condition."
    effective_gate = gate or level == 3 or (level == 2 and priority == "Critical")
    CASES.append({
        "Case ID": cid, "Level": f"Level {level}", "Module": module, "Submodule": "",
        "Scenario": title, "Test Type": test_type, "Role": role, "Priority": priority,
        "Release Gate": "Yes" if effective_gate else "No", "Gap Covered": gap,
        "Preconditions": precondition, "Test Data": test_data, "Steps": steps,
        "Expected Result": expected, "Status": "Not Run", "Actual Result": "",
        "Defect ID": "", "Defect Severity": "", "Defect Status": "",
        "Tester": "", "Environment": "QA", "Build": "", "Execution Date": "",
        "Evidence Link": "", "Retest Status": "Not Required", "Comments": "",
    })


def add_group(level, module, role, specs, default_type="Functional", default_priority="Medium"):
    for spec in specs:
        title, action, expected = spec[:3]
        opts = spec[3] if len(spec) > 3 else {}
        add_case(level, module, title, action, expected, opts.get("role", role),
                 opts.get("type", default_type), opts.get("priority", default_priority),
                 opts.get("gap", ""), opts.get("gate", False),
                 opts.get("precondition", "Required master data and user exist."),
                 opts.get("data", "Use dedicated non-production test data."))


def build_catalog():
    # LEVEL 1 - core smoke and complete happy path, including the steps missed in the original flow.
    l1 = [
        ("Platform", "Provision a lab tenant", "Create a lab with an active plan and required modules.", "The tenant is created, isolated and can open its login page.", "Developer Admin"),
        ("Settings/RBAC", "Create the test users", "Create Admin, Front Desk, Cashier, Phlebotomist, Technician, Pathologist, Accounts and Inventory users.", "Every user can sign in and sees only the expected menus.", "Admin"),
        ("Test Master", "Create test masters", "Create a category, test, parameters, reference ranges, sample type, price and package.", "The test and package are active and available for billing.", "Admin"),
        ("Inventory", "Load opening stock", "Create required inventory masters and load stock with batch and expiry.", "On-hand stock, batch, UOM and expiry are correct.", "Inventory Manager"),
        ("Doctors", "Create and activate a regular doctor", "Create a doctor and complete invitation activation.", "The doctor can sign in to the correct tenant portal.", "Admin"),
        ("Patients", "Search before patient registration", "Search by phone and patient ID before creating a patient.", "The system shows existing matches and supports safe new registration.", "Front Desk Receptionist"),
        ("Patients", "Register a new patient", "Enter valid demographics and save the patient.", "A unique patient ID and portal account are created.", "Front Desk Receptionist"),
        ("Billing", "Create a test order and bill", "Select patient, doctor, individual test, package and priority; create the bill.", "The bill total is correct and linked sample records are created.", "Front Desk Receptionist"),
        ("Billing", "Record a partial payment", "Pay less than the bill balance and issue a receipt.", "Status becomes Partial and the remaining balance is correct.", "Billing Cashier"),
        ("Samples", "Print or open sample barcode", "Open the generated sample and barcode.", "Sample ID, barcode, patient, bill and required sample type match.", "Phlebotomist"),
        ("Samples", "Collect the sample", "Verify two patient identifiers, record collection/received details and mark Collected.", "Collection time, collector and custody event are saved.", "Phlebotomist"),
        ("Samples", "Move sample to processing", "Open the collected sample and start processing.", "Sample and bill-item statuses move forward in the correct order.", "Lab Technician"),
        ("Inventory", "Reserve or monitor test inventory", "Check required item availability during sample processing.", "Reserved/on-hand stock and warnings follow the implemented rule.", "Lab Technician"),
        ("Reports", "Enter valid results", "Enter all required numeric and text results and complete the sample.", "Values, units and flags are correct; a Draft report exists.", "Lab Technician"),
        ("Reports", "Review the report", "Check patient, sample, tests, results, ranges and remarks; mark Reviewed.", "The report becomes Reviewed and records the reviewer and time.", "Pathologist"),
        ("Reports", "Approve the report", "Approve the Reviewed report.", "The report becomes Approved and records the approver and time.", "Pathologist"),
        ("Reports", "Release the report", "Release the Approved report.", "The report becomes Released and is available to authorized portals.", "Pathologist"),
        ("Notifications", "Check report release notification", "Refresh notifications after report release.", "Authorized users receive the correct report-release notification where implemented.", "Lab Manager"),
        ("Doctor Portal", "View the referred patient's report", "Sign in as the linked doctor and open the referred patient.", "Only owned referrals and Released reports are visible.", "Doctor Regular"),
        ("Patient Portal", "Activate patient access", "Issue QR/PIN, activate with DOB and set the portal PIN.", "The patient session is created and old one-time credentials cannot be reused.", "Patient"),
        ("Patient Portal", "View patient bill and report", "Open Bills, Receipts and Reports.", "Only the patient's own financial data and Released reports are visible.", "Patient"),
        ("Billing", "Complete payment", "Pay the remaining balance and issue the final receipt.", "Bill status becomes Paid and balance becomes zero.", "Billing Cashier"),
        ("Accounts", "Reconcile invoice and payments", "Compare bill, receipts, journal, ledger and doctor commission.", "All amounts balance and record IDs link correctly.", "Accounts Manager"),
        ("Audit", "Check core audit evidence", "Review audit records for login, billing, sample and report release.", "Available audit records identify actor, action, time and object; missing coverage is logged as a gap.", "Admin"),
        ("Security", "Check tenant isolation", "Try to open a record from another tenant using its ID.", "The record is denied or not found and no data is disclosed.", "Admin"),
        ("Workflow", "Run the complete second patient flow", "Repeat patient-to-report flow with a package, urgent priority and full payment.", "The complete workflow passes without stale or mismatched links.", "QA Tester"),
    ]
    for module, title, action, expected, role in l1:
        add_case(1, module, title, action, expected, role, "Smoke", "Critical" if module in ("Security", "Reports", "Billing") else "High", gate=True,
                 gap="Included missed flow step" if module in ("Platform", "Inventory", "Notifications", "Audit", "Security") else "")

    # LEVEL 2 - detailed functional, integration, negative, boundary and gap testing.
    add_group(2, "Authentication & Session", "Tenant User", [
        ("Valid login", "Enter valid tenant credentials.", "Dashboard opens and a secure tenant session is created."),
        ("Invalid password", "Enter a valid user ID with a wrong password.", "Login fails without revealing sensitive account details.", {"type":"Negative","priority":"High"}),
        ("Inactive user", "Disable a user and attempt login.", "Login is denied and the existing session is no longer useful.", {"type":"Security","priority":"Critical"}),
        ("Temporary lockout", "Submit repeated wrong passwords.", "Rate limit or temporary lockout is applied."),
        ("Forgot password", "Request reset for registered and unregistered emails.", "Response does not expose whether an account exists."),
        ("Reset token reuse", "Use a reset credential twice.", "The second use is denied.", {"type":"Security","priority":"High"}),
        ("Password policy", "Try weak, mismatched and valid new passwords.", "Only the valid password is accepted."),
        ("Logout", "Logout, then refresh and use browser Back.", "Protected data is not available after logout."),
        ("Session expiry", "Use an expired session on UI and API.", "The user is asked to sign in again."),
        ("Wrong tenant login", "Use valid credentials on another lab tenant.", "Access is denied without data leakage.", {"type":"Security","priority":"Critical"}),
    ])

    add_group(2, "RBAC & Settings", "Admin", [
        ("Create custom role", "Create a limited role with selected permissions.", "Only selected permissions are effective."),
        ("Permission dependency", "Select a permission that requires another permission.", "The dependency is enforced or clearly added."),
        ("Dangerous permission grant", "Grant refund, release or user-management permission.", "The grant follows the approved control and is auditable.", {"gap":"Dangerous-permission approval flow","priority":"Critical"}),
        ("Assigned role deletion", "Try to delete a role assigned to a user.", "Deletion is blocked."),
        ("Role change during session", "Change a signed-in user's role.", "New permissions apply after the documented refresh/re-login rule."),
        ("Disabled module UI", "Disable a module for the tenant.", "Its navigation and actions disappear."),
        ("Disabled module API", "Call the disabled module API directly.", "The server denies the action."),
        ("Technician collection denial", "Use Lab Technician to collect a sample.", "The default role is denied."),
        ("Technician release denial", "Use Lab Technician to release a report.", "The default role is denied."),
        ("Front Desk broad permissions", "Check test-delete, sample-update and report-edit actions.", "Behavior matches the approved Front Desk policy.", {"gap":"Front Desk template is broader than its description","priority":"Critical"}),
        ("Role change audit", "Create and update a role.", "Actor, old/new permissions and time are auditable.", {"gap":"Complete permission-change audit","priority":"High"}),
        ("Inactive role", "Deactivate a role and try to use it.", "The role cannot grant access."),
    ])

    add_group(2, "Patients & Visits", "Front Desk Receptionist", [
        ("Required fields", "Submit each required field as blank.", "Clear field validation blocks save."),
        ("Phone and DOB validation", "Enter invalid phone and future DOB.", "Invalid values are blocked."),
        ("Duplicate warning", "Register matching name and phone.", "The configured duplicate warning/block appears."),
        ("Forced duplicate", "Confirm a permitted duplicate.", "Creation requires explicit confirmation and remains traceable."),
        ("Patient search", "Search by name, phone, patient ID and barcode.", "Correct records appear and permissions are respected."),
        ("Special-character search", "Search with + . * ( ).", "Search is escaped and does not crash.", {"type":"Negative"}),
        ("Edit patient", "Change allowed demographics and refresh.", "Changes persist and linked records remain intact."),
        ("Repeat visit", "Create a new visit for an existing patient.", "History keeps both visits under one patient."),
        ("Delete linked patient", "Try to delete a patient with bills/reports.", "Deletion is blocked or follows the approved safe policy.", {"gap":"Deletion protection for linked patients","priority":"Critical"}),
        ("Soft delete and restore", "Delete and restore a patient.", "History and links are preserved.", {"gap":"Patient soft-delete/restore is pending"}),
        ("Patient edit audit", "Change identity-sensitive fields.", "Old/new values, actor and time are auditable.", {"gap":"Patient audit coverage"}),
        ("Pagination", "Use first, middle and last patient pages.", "Counts, sorting and navigation are correct."),
    ])

    add_group(2, "Doctors & Commission", "Admin", [
        ("Doctor required fields", "Submit invalid or missing doctor fields.", "Save is blocked with clear validation."),
        ("Duplicate email", "Create two doctors with the same email.", "Duplicate doctor/portal identity is blocked."),
        ("Invitation activation", "Use a valid doctor invitation OTP.", "The account activates once and can sign in."),
        ("Expired invitation", "Use an expired invitation.", "Activation is denied."),
        ("Resend invitation", "Resend and try old and new OTPs.", "Only the latest OTP works."),
        ("Inactive doctor billing", "Select an inactive referral doctor.", "New billing is blocked."),
        ("Inactive doctor login", "Deactivate a doctor with an active user.", "Portal login is denied."),
        ("Historical commission", "Change doctor rate after an old bill.", "Old commission does not change."),
        ("Commission eligibility", "Compare unpaid, partial and paid referrals.", "Estimated and earned commission follow policy."),
        ("Payout approval", "Create and release a payout.", "Only authorized staff can approve and the action is auditable.", {"gap":"Commission approval workflow","priority":"High"}),
        ("Payout history", "Open completed payout history.", "Doctor-wise statement and references are complete.", {"gap":"Detailed payout statement/history"}),
        ("Doctor delete protection", "Try to delete a doctor with referrals.", "Deletion is blocked or safely archived.", {"gap":"Doctor soft delete and linked-record protection"}),
    ])

    add_group(2, "Test Master", "Lab Manager", [
        ("Category duplicate", "Create a duplicate category.", "Duplicate is blocked."),
        ("Test code duplicate", "Create a duplicate test code.", "Duplicate is blocked."),
        ("Parameter validation", "Create required numeric/text parameters with units.", "Definitions save correctly."),
        ("Reference ranges", "Configure male/female and boundary ranges.", "Flags are correct at, below and above boundaries."),
        ("Package expansion", "Create a package with multiple tests.", "All active tests expand exactly once."),
        ("Package invalid member", "Add inactive/duplicate tests to a package.", "Invalid package setup is blocked.", {"gap":"Stronger package validation"}),
        ("Price permission", "Edit price without the required permission.", "Server and UI deny the edit."),
        ("Deactivate test", "Deactivate a test after old bills exist.", "Old snapshots remain; new billing cannot select it."),
        ("Restore test", "Restore a deactivated test.", "Availability follows the approved restore policy.", {"gap":"Test deactivate/restore workflow"}),
        ("Age-specific ranges", "Configure/test age-dependent references.", "Results use the correct range or the gap is recorded.", {"gap":"Advanced age/context reference ranges"}),
        ("Inventory requirement", "Link item and quantity/UOM to a test.", "The requirement is available to sample processing."),
        ("Master change audit", "Edit price, range and parameters.", "Old/new values and actor are auditable.", {"gap":"Complete test-master audit"}),
    ])

    add_group(2, "Billing & Payments", "Billing Cashier", [
        ("Manual bill", "Create a bill with one test.", "ID, line, price and total are correct."),
        ("Package bill", "Bill an individual test plus a package.", "No duplicated lines and correct total."),
        ("Discount amount", "Apply a permitted fixed discount.", "Net total is correct."),
        ("Discount limit", "Apply an excessive or unauthorized discount.", "Action is blocked and audited."),
        ("Tax calculation", "Apply configured tax.", "Tax rounding and final total are correct."),
        ("Zero/negative amount", "Try invalid price, discount or payment values.", "Invalid values are blocked."),
        ("Full payment", "Pay the exact remaining amount.", "Status becomes Paid and receipt/journal are created."),
        ("Partial payment", "Pay less than the balance.", "Status becomes Partial and due is correct."),
        ("Multiple modes", "Pay using cash, card, UPI, cheque and corporate credit.", "Mode breakdown and history are correct."),
        ("Overpayment", "Pay more than the balance.", "Payment is blocked with no partial posting."),
        ("Duplicate submission", "Double-click or repeat the same payment request.", "Only one payment/receipt is created.", {"type":"Integration","priority":"Critical"}),
        ("Cancellation reason", "Cancel a bill without and with a reason.", "Reason is required and downstream behavior follows policy."),
        ("Refund", "Refund a paid/cancelled bill.", "Authorized refund reverses amounts and records history.", {"gap":"Refund-flow polish","priority":"Critical"}),
        ("Daily close", "Close the day and try late changes.", "Close/reopen policy and totals are enforced.", {"gap":"Daily close and reconciliation","priority":"High"}),
        ("Corporate credit", "Bill an active corporate account.", "Receivable and statement are correct."),
        ("Credit limit", "Exceed corporate credit limit.", "Warning/block follows policy.", {"gap":"Corporate credit-limit alerts"}),
        ("Invoice/receipt print", "Generate invoice and all receipts.", "Patient, lab, totals, IDs and branding are correct."),
        ("Settlement audit", "Collect, cancel, discount and refund.", "Every sensitive action is auditable.", {"gap":"Complete billing audit"}),
    ])

    add_group(2, "Samples", "Phlebotomist", [
        ("Automatic sample creation", "Create a multi-test bill.", "Correct samples/investigations are linked."),
        ("Unique sample ID", "Create many samples concurrently.", "Every sample ID and barcode is unique."),
        ("Collection time validation", "Enter future or pre-DOB collection time.", "Invalid time is blocked."),
        ("Received time order", "Enter received time before collection.", "Invalid order is blocked."),
        ("Valid status sequence", "Move Registered to Collected to Processing to Completed.", "Each transition persists with custody history."),
        ("Invalid status jump", "Jump Registered directly to Completed/Released.", "Transition is blocked."),
        ("Reject without reason", "Reject an active sample without a reason.", "Save is blocked."),
        ("Reject with reason", "Reject from each allowed active state.", "Reason and custody event are retained."),
        ("Recollection", "Create a replacement after rejection.", "Old rejected sample remains traceable and the new sample links correctly.", {"gap":"Controlled recollection workflow","priority":"Critical"}),
        ("Barcode print and scan", "Print and scan the label in collection and processing.", "Scan resolves the correct sample with readable label.", {"gap":"Barcode printing/scanning polish","priority":"High"}),
        ("Patient/sample mismatch", "Use another patient's barcode.", "Processing is blocked.", {"priority":"Critical"}),
        ("Batch/custody", "Transfer a sample between users/locations.", "Complete chain of custody is visible."),
        ("Archive and restore", "Archive and restore a completed sample.", "Policy is enforced without losing history.", {"gap":"Sample archive/restore controls"}),
        ("Rejection approval", "Reject a high-impact sample and request approval.", "Approval follows policy or is recorded as a gap.", {"gap":"Stronger rejection approval"}),
        ("Status vocabulary integration", "Create a report from Processing and Completed samples.", "Report creation accepts the model's real statuses.", {"gap":"API/model sample-status mismatch","priority":"Critical"}),
        ("Sample audit", "Collect, reject and complete samples.", "All sensitive sample changes are auditable.", {"gap":"Sample audit coverage"}),
    ])

    add_group(2, "Reports", "Pathologist", [
        ("Required result", "Leave a required result empty.", "Review is blocked."),
        ("Invalid numeric result", "Enter text or exponential notation in a numeric field.", "Invalid input is blocked."),
        ("Boundary flags", "Enter values at and outside reference limits.", "Normal/low/high flags are correct."),
        ("Draft edit", "Edit a Draft report.", "Changes save and version behavior is correct."),
        ("Reviewed edit", "Try to edit a Reviewed report as a typist.", "Unauthorized/status-invalid edit is blocked."),
        ("Status sequence", "Move Draft to Reviewed to Approved to Released.", "Only the correct sequence works."),
        ("Skip approval", "Release Draft or Reviewed report.", "Release is blocked."),
        ("Wrong role release", "Release as Technician or Front Desk without permission.", "Server and UI deny release."),
        ("Previous versions", "Create a new report version.", "Prior values and metadata remain available to staff."),
        ("Released-only portal", "Query every report status from both portals.", "Only Released is returned."),
        ("Print permission", "Print as allowed and disallowed roles.", "Permission is enforced."),
        ("Branded PDF", "Download final report PDF.", "Brand, patient data, tests, units, ranges and pagination are correct.", {"gap":"Final branded PDF polish"}),
        ("Digital signature", "Release a signed report and verify signature display.", "Signature is authentic and shown correctly or the gap is blocked.", {"gap":"Digital signature","priority":"High"}),
        ("Post-release correction", "Correct a Released report with a reason and approval.", "Original remains; amendment is approved and portals show correct version.", {"gap":"Correction approval workflow","priority":"Critical"}),
        ("Critical result", "Enter a configured critical result.", "Alert/acknowledgement follows policy.", {"gap":"Critical-value workflow decision","priority":"Critical"}),
        ("Release audit", "Review, approve, release and amend.", "Actor, time, status and version are auditable.", {"gap":"Complete report audit","priority":"Critical"}),
    ])

    add_group(2, "Inventory", "Inventory Manager", [
        ("Item master validation", "Create items with invalid/duplicate codes and valid data.", "Only valid unique item saves."),
        ("UOM conversion", "Receive and consume stock in different linked UOMs.", "Base quantities convert correctly."),
        ("Purchase order lifecycle", "Create, update and receive a purchase order.", "Ordered and received quantities are correct."),
        ("Batch receipt", "Receive two batches with different expiry dates.", "Both batches and totals are correct."),
        ("Stock movement", "Post receipt, issue, adjustment and transfer.", "On-hand and movement history reconcile."),
        ("Negative stock", "Issue more than available.", "Negative stock is blocked or controlled by approved override."),
        ("Low stock", "Move stock below reorder level.", "Low-stock warning/notification appears.", {"gap":"Complete low-stock alert workflow"}),
        ("Expired stock", "Try to reserve/use an expired batch.", "Expired stock is blocked."),
        ("Near-expiry", "Create a near-expiry batch.", "Expiry dashboard/alert appears.", {"gap":"Expiry dashboard"}),
        ("Automatic consumption", "Complete a test linked to inventory.", "Configured quantity is consumed once or the gap blocks release.", {"gap":"Automatic inventory consumption","priority":"Critical"}),
        ("Insufficient stock", "Start a test with insufficient required stock.", "Warning/block follows the approved policy."),
        ("Reservation release", "Reject/cancel a sample after reservation.", "Reserved stock is released correctly."),
        ("Wastage", "Record damaged/expired wastage.", "Wastage changes stock and remains traceable.", {"gap":"Wastage tracking"}),
        ("Purchase approval", "Approve/reject a purchase order.", "Only authorized users can approve.", {"gap":"Purchase approval workflow"}),
        ("Import/export", "Import valid/invalid inventory and export the result.", "Valid rows load, errors are clear and export reconciles."),
        ("Inventory audit", "Edit item and post movement.", "Old/new values and movement actor are auditable.", {"gap":"Inventory audit coverage"}),
    ])

    add_group(2, "Accounts, Corporate & Expenses", "Accounts Manager", [
        ("Invoice journal", "Create a bill.", "Receivable/revenue journal is balanced and linked."),
        ("Payment journal", "Collect a payment.", "Cash/bank and receivable entries are balanced."),
        ("Refund journal", "Refund a payment.", "Reversal is balanced and linked."),
        ("Journal balance", "Review generated journals.", "Total debit equals total credit."),
        ("Manual journal permission", "Create a manual journal as allowed/disallowed users.", "Permission and balanced-entry rules are enforced."),
        ("Ledger filters", "Filter by account/date/type and paginate.", "Rows and totals match source entries."),
        ("Daily collection", "Compare receipts with daily collection report.", "Totals and payment modes reconcile."),
        ("Outstanding", "Compare unpaid/partial bills with outstanding report.", "Patient and corporate dues match."),
        ("P&L", "Compare revenue and expense source records.", "P&L totals and date ranges are correct."),
        ("Expense validation", "Create valid and invalid expenses.", "Only valid authorized expense saves."),
        ("Expense approval", "Submit and approve/reject an expense.", "Approval history is complete.", {"gap":"Expense approval workflow"}),
        ("Expense attachment", "Attach supporting evidence.", "Attachment is safe and retrievable or the gap is recorded.", {"gap":"Expense attachments"}),
        ("Recurring expense", "Create and generate a recurring expense.", "No missed or duplicate entries.", {"gap":"Recurring expenses"}),
        ("Corporate statement", "Generate a statement for a date range.", "Opening, bills, payments and closing balance reconcile."),
        ("Corporate aging", "Age outstanding corporate invoices.", "Aging buckets are correct.", {"gap":"Corporate aging report"}),
        ("Daily close/reopen", "Close, attempt late entry and authorized reopen.", "Policy, reason and audit are enforced.", {"gap":"Daily close/reconciliation"}),
        ("Financial export", "Export ledger, P&L and statements.", "Export matches screen totals.", {"gap":"Advanced export formats"}),
        ("Accounts audit", "Post, edit, refund and close.", "Complete financial audit history exists.", {"gap":"Accounts audit coverage","priority":"Critical"}),
    ])

    add_group(2, "Doctor & Patient Portals", "Portal User", [
        ("Doctor owned referral", "Open a patient referred to the signed-in doctor.", "Owned patient data and Released reports are visible."),
        ("Doctor non-owned referral", "Request another doctor's patient ID.", "Not found/denied without existence disclosure.", {"type":"Security","priority":"Critical"}),
        ("Doctor draft report", "Request Draft/Reviewed/Approved report IDs.", "All are hidden."),
        ("Doctor profile edit", "Edit permitted and restricted doctor fields.", "Only allowed fields change."),
        ("Investor access", "Open investor-only analytics, patients, bills and accounts.", "Only approved Investor role/plan can access.", {"gap":"Investor privacy and scope decision","priority":"Critical"}),
        ("Doctor-created request", "Create patient/test request from doctor portal.", "Request enters the correct lab work queue without bypassing staff controls."),
        ("Patient activation", "Activate QR/PIN with correct token, access PIN and DOB.", "One-time credentials are cleared and session is created."),
        ("Wrong patient DOB/PIN", "Use wrong DOB or PIN repeatedly.", "Access is denied and lockout applies."),
        ("Patient cross-record", "Request another patient's bill/report ID.", "Denied with no data leakage.", {"type":"Security","priority":"Critical"}),
        ("Reissue patient slip", "Issue a new slip after active login.", "Old credentials and sessions are invalidated."),
        ("Patient data minimization", "Inspect patient portal response.", "No commission, journal, staff, draft or version-history fields appear.", {"type":"Privacy","priority":"Critical"}),
        ("No-store/no-index", "Inspect patient portal response headers and page metadata.", "Private no-store and no-index controls are present."),
        ("Portal print", "Print a Released report from patient and doctor portals.", "Correct report prints without hidden internal data."),
        ("Portal audit", "Login and open a report.", "Portal access events are auditable."),
    ])

    add_group(2, "Dashboard, Search, Notifications & Analytics", "Lab Manager", [
        ("Role-specific dashboard", "Compare cards for several roles.", "Each role sees permitted, tenant-correct metrics."),
        ("Empty dashboard", "Use a tenant with no data.", "Zero/empty state is clear and stable."),
        ("Global search permission", "Search as limited roles.", "Results include only permitted modules/records."),
        ("Search navigation", "Open patient, sample, report and bill search results.", "Each link opens the correct record."),
        ("Low-stock notification", "Trigger low stock.", "Correct users receive the notification."),
        ("Pending report notification", "Create overdue/pending work.", "Configured notification appears."),
        ("Read/unread state", "Mark notification read and reload.", "State persists.", {"gap":"Richer read/unread behavior"}),
        ("Notification preferences", "Disable a notification category.", "Preference is respected.", {"gap":"User-specific notification preferences"}),
        ("Analytics date filter", "Compare analytics with source bills in a date range.", "Revenue and counts match."),
        ("Doctor analytics privacy", "Compare Regular and Investor analytics access.", "Financial/commission data follows approved permissions."),
        ("Pending workflow analytics", "Compare pending samples/reports with operational lists.", "Counts reconcile or the gap is recorded.", {"gap":"Deeper pending-work analytics"}),
        ("Analytics export", "Export an analytics report.", "Export matches filtered screen data.", {"gap":"Exportable analytics reports"}),
    ])

    add_group(2, "Developer, Subscription & Tenant", "Developer Admin", [
        ("Tenant provisioning rollback", "Force failure during onboarding.", "No half-created tenant or admin remains.", {"type":"Integration","priority":"Critical"}),
        ("Tenant status", "Disable and re-enable a lab.", "Tenant login follows status immediately."),
        ("Module plan", "Change subscription plan/modules.", "Navigation and APIs follow the new plan."),
        ("Runtime quota", "Exceed a configured patient/report/user quota.", "The server enforces quota consistently.", {"gap":"Complete runtime subscription enforcement","priority":"Critical"}),
        ("Upgrade request", "Submit and confirm a subscription upgrade.", "Plan changes once with traceable history."),
        ("Addon request", "Submit and confirm an addon.", "Addon applies once and usage limits update."),
        ("Tenant archive/delete", "Archive/delete a lab with data.", "Lifecycle policy protects or recovers data."),
        ("Developer-only route", "Open developer routes as tenant staff.", "Access is denied."),
        ("Cross-tenant search", "Search tenant IDs from another tenant session.", "No other tenant data appears.", {"type":"Security","priority":"Critical"}),
        ("Developer audit", "Create, disable and change a tenant plan.", "Developer actions are auditable.", {"gap":"Developer audit events"}),
    ])

    # LEVEL 3 - release certification / market readiness.
    add_group(3, "Security", "Security Tester", [
        ("Broken object authorization", "Change patient, bill, sample, report, doctor and account IDs in requests.", "Every unauthorized object request is denied without data leakage."),
        ("Cross-tenant authorization", "Replay valid requests with IDs from another tenant.", "Tenant isolation holds for every tested API."),
        ("Role escalation", "Modify client payloads/cookies and call admin/refund/release APIs.", "Server-side authorization prevents escalation."),
        ("Session fixation", "Compare session before and after login/activation.", "Authentication rotates or safely establishes the session."),
        ("Cookie security", "Inspect all staff, doctor and patient cookies.", "Secure, HttpOnly and appropriate SameSite/scope are applied."),
        ("CSRF-sensitive actions", "Attempt cross-site state-changing requests.", "Requests are blocked by the implemented session/CSRF design."),
        ("Injection inputs", "Submit SQL/NoSQL/operator/script payloads in search and forms.", "Input is validated/escaped and data is not exposed or changed."),
        ("Stored XSS", "Save script/HTML payloads in names, notes and remarks.", "Payload renders as safe text or is sanitized."),
        ("File upload validation", "Upload oversized, wrong-type and disguised files/images.", "Unsafe files are rejected and allowed files are isolated."),
        ("OTP brute force", "Repeat doctor/patient/reset OTP guesses.", "Rate limit, expiry and lockout prevent brute force."),
        ("Credential reuse", "Reuse activation/reset/QR tokens after success and reissue.", "All old single-use credentials fail."),
        ("Error information leakage", "Trigger invalid IDs, server errors and unauthorized requests.", "Responses contain no stack, DB, secret or tenant details."),
        ("Security headers", "Inspect login, dashboard and portal headers.", "Required transport/content framing/sniffing controls meet policy."),
        ("Secrets exposure", "Inspect client bundles, logs and error output.", "No credentials, SMTP, database or signing secrets are exposed."),
        ("Dependency vulnerability gate", "Run the approved dependency/security scan.", "No unaccepted Critical/High finding remains."),
        ("Privilege matrix certification", "Execute every dangerous permission against allowed and denied roles.", "UI and API match the signed-off matrix."),
    ], default_type="Security", default_priority="Critical")

    add_group(3, "Privacy & Medical Data", "Privacy Tester", [
        ("Minimum necessary patient response", "Inspect all patient portal/API fields.", "Only approved patient-safe fields are returned."),
        ("Minimum necessary doctor response", "Inspect doctor portal/API fields.", "Only owned clinical/referral/approved commission data is returned."),
        ("URL and log privacy", "Inspect URLs, client/server logs and analytics events.", "Sensitive medical data and tokens are not logged or placed in URLs."),
        ("Cache privacy", "Use Back, shared browser profile and proxy cache after logout.", "Private pages are not served from unsafe cache."),
        ("Search engine privacy", "Inspect robots/meta behavior for patient pages.", "Patient pages are not indexed."),
        ("Printout privacy", "Review QR slips, invoices and reports.", "Only necessary data is printed and confidential slips are clear."),
        ("Data export authorization", "Export reports/accounts/inventory as several roles.", "Exports contain only authorized tenant data."),
        ("Data retention", "Age test audit, portal and medical records past configured periods.", "Retention/deletion follows approved policy."),
        ("Support access", "Use support/developer access to a tenant.", "Access is authorized, time-bound and auditable."),
        ("Privacy incident evidence", "Trace one patient record across access logs.", "The organization can identify who accessed or changed it."),
    ], default_type="Privacy", default_priority="Critical")

    add_group(3, "Performance & Scale", "Performance Tester", [
        ("Concurrent login", "Run expected peak staff/doctor/patient logins.", "Login remains within agreed response time and error rate."),
        ("Patient search scale", "Search a production-sized patient dataset.", "Results return within the agreed SLA and remain correct."),
        ("Billing concurrency", "Create bills and payments concurrently.", "No duplicate IDs, receipts or lost payments occur."),
        ("Sample concurrency", "Create/collect/process many samples concurrently.", "IDs/statuses/custody remain correct."),
        ("Report entry concurrency", "Enter and release multiple reports concurrently.", "No overwritten results or invalid transitions occur."),
        ("Dashboard load", "Open dashboards during peak transactions.", "Dashboard does not degrade core workflow beyond SLA."),
        ("Large pagination", "Use first/middle/last pages for large datasets.", "Stable ordering, counts and response time meet SLA."),
        ("Large report PDF", "Generate reports with many tests/parameters.", "PDF completes within SLA with correct layout."),
        ("Export volume", "Export production-sized ledger/inventory datasets.", "Export completes without timeout or memory failure."),
        ("Notification burst", "Release many reports and trigger stock alerts.", "Notifications do not duplicate or block core work."),
        ("Long-running soak", "Run realistic transactions for the agreed soak duration.", "Memory, connection and error rates remain stable."),
        ("Recovery after load", "Stop the load and continue normal use.", "The system returns to normal without manual restart."),
    ], default_type="Performance", default_priority="High")

    add_group(3, "Reliability, Backup & Recovery", "Release Tester", [
        ("Transactional rollback", "Force failure during patient+bill+sample creation.", "No partial linked records remain."),
        ("Payment rollback", "Force failure after payment starts.", "No mismatched bill, receipt or journal remains."),
        ("Report release rollback", "Force failure during release/notification.", "Report status and downstream events remain consistent."),
        ("Network interruption", "Disconnect during save/payment/result entry.", "User sees a safe retry state with no duplicate transaction."),
        ("Browser refresh during submit", "Refresh during sensitive submissions.", "No duplicate or corrupt records are created."),
        ("Database reconnect", "Interrupt database connectivity in a controlled environment.", "System fails safely and recovers without cross-tenant impact."),
        ("Backup creation", "Run the production backup procedure.", "Backup completes, is encrypted and produces verifiable evidence."),
        ("Restore drill", "Restore backup into an isolated environment.", "Patients, bills, samples, reports, accounts and tenant links reconcile."),
        ("Point-in-time recovery", "Recover to an agreed point after test transactions.", "Recovery meets RPO and data integrity expectations."),
        ("Disaster recovery", "Execute the approved DR runbook.", "RTO/RPO, ownership and communications are proven."),
        ("Job retry/idempotency", "Retry email, notification and background operations.", "Retries are controlled and do not duplicate business records."),
        ("Graceful maintenance", "Enable maintenance/deploy while users are active.", "Users receive safe behavior and no data corruption occurs."),
    ], default_type="Recovery", default_priority="Critical")

    add_group(3, "Compatibility & Accessibility", "Compatibility Tester", [
        ("Chrome current", "Run the full Level 1 flow on current Chrome.", "All screens and downloads work."),
        ("Edge current", "Run the full Level 1 flow on current Edge.", "All screens and downloads work."),
        ("Firefox current", "Run the full Level 1 flow on current Firefox.", "All screens and downloads work."),
        ("Mobile patient portal", "Run activation, bills and reports at 390px width.", "No blocked controls, clipping or unsafe horizontal overflow."),
        ("Tablet staff use", "Run reception/sample flow at 768px width.", "Primary actions and tables remain usable."),
        ("Desktop resolutions", "Check 1366x768 and 1920x1080.", "Layout, dialogs and tables remain readable."),
        ("Keyboard navigation", "Complete critical forms without a mouse.", "Focus order, controls and dialogs are usable."),
        ("Screen reader labels", "Inspect login, patient, billing and portal forms.", "Controls and errors have meaningful accessible names."),
        ("Color and contrast", "Review status, error, flag and dashboard colors.", "Meaning is not color-only and contrast meets policy."),
        ("Zoom/text scaling", "Use 200% browser zoom.", "Critical workflows remain operable without hidden content."),
        ("Print compatibility", "Print invoice, receipt, barcode and report from supported browsers.", "Content is complete, readable and not clipped."),
        ("Locale/timezone", "Use supported locale/timezone variants.", "Dates, times and numbers remain accurate and unambiguous."),
    ], default_type="Compatibility", default_priority="High")

    add_group(3, "Clinical & Financial Integrity", "Senior QA", [
        ("Patient identity chain", "Trace patient ID from registration through report and portal.", "No identity mismatch occurs."),
        ("Order-sample-report chain", "Trace each billed item through sample and report.", "Every item has correct one-to-one/many relationship and status."),
        ("Reference range certification", "Verify representative sex/age/boundary results.", "Displayed ranges and flags match the signed-off master."),
        ("Report release segregation", "Attempt entry, approval and release with all roles.", "Signed-off segregation is enforced."),
        ("Critical-result handling", "Trigger critical results end to end.", "Notification/acknowledgement policy is proven or release is blocked."),
        ("Bill-to-ledger reconciliation", "Reconcile all UAT bills/payments/refunds.", "Invoice, receipt, journal, ledger and P&L totals agree."),
        ("Commission reconciliation", "Reconcile regular/investor referral commissions and payouts.", "Historical rate, eligibility and payout references agree."),
        ("Corporate reconciliation", "Reconcile credit bills, payments, statements and aging.", "Opening, movement and closing balances agree."),
        ("Inventory reconciliation", "Reconcile opening, receipts, consumption, rejects and wastage.", "Physical/system balance and reservations agree."),
        ("Audit completeness", "Sample sensitive actions from every module.", "Required actor/action/time/object/before-after/reason evidence exists."),
        ("Status synchronization", "Compare bill, item, sample and report statuses after all paths.", "No impossible or stale combination remains."),
        ("No orphan records", "Query for unlinked bills, samples, reports, receipts and journals.", "No unexplained orphan or duplicate record exists."),
    ], default_type="Integrity", default_priority="Critical")

    add_group(3, "Production Operations", "Release Manager", [
        ("Production configuration", "Validate domains, HTTPS, SMTP, storage, database and secrets.", "All production settings are present and no development fallback is active."),
        ("Tenant domain routing", "Open valid/invalid tenant domains.", "Correct tenant resolution and safe not-found behavior occur."),
        ("Monitoring", "Trigger application, API, DB, email and storage failures.", "Monitoring detects them with actionable context."),
        ("Alert routing", "Trigger Critical and High operational alerts.", "The correct on-call owner receives alerts within SLA."),
        ("Log correlation", "Trace a transaction across UI/API/database/audit logs.", "A correlation path exists without exposing medical data."),
        ("Health checks", "Check production health/readiness endpoints during dependency failure.", "Health state is accurate for routing and operations."),
        ("Deployment rollback", "Deploy a test release and execute rollback.", "Rollback is documented, fast and data-safe."),
        ("Database migration", "Run migration on a production-like copy twice.", "Migration is idempotent or safely guarded and preserves data."),
        ("Support runbook", "Exercise login, email, payment, report and tenant-support procedures.", "Support can diagnose and escalate with defined ownership."),
        ("Incident runbook", "Simulate patient-data or report-release incident.", "Containment, evidence, notification and recovery steps are executable."),
        ("Capacity thresholds", "Compare load results with infrastructure limits.", "Documented capacity and scale-up thresholds exist."),
        ("License/subscription enforcement", "Validate plan and quota behavior in production-like setup.", "Commercial entitlements are enforced without blocking entitled labs."),
    ], default_type="Operational", default_priority="Critical")

    add_group(3, "Release Certification & Sales UAT", "Business UAT Team", [
        ("New lab onboarding", "Onboard a clean customer-like lab without developer fixes.", "Admin can configure and start operations using documented steps."),
        ("Front-desk UAT", "Run registration, billing, receipt and portal issue as real staff.", "Workflow is understandable and completes within agreed time."),
        ("Lab UAT", "Run collection, processing, rejection and report entry.", "Work queues and hand-offs are clear and reliable."),
        ("Pathologist UAT", "Review, approve, release, print and correct reports.", "Clinical controls and outputs meet sign-off expectations."),
        ("Accounts UAT", "Run collection, refund, close, reconciliation and reporting.", "Financial controls and reports meet sign-off expectations."),
        ("Inventory UAT", "Receive, issue, consume, alert and reconcile stock.", "Inventory is usable for daily lab operations."),
        ("Doctor UAT", "Activate, view referrals/reports/commission and create a request.", "Doctor experience and privacy meet sign-off expectations."),
        ("Patient UAT", "Activate, log in, view and print bills/reports on mobile.", "Patient experience is clear, private and usable."),
        ("Training validation", "Give role guides to new users and observe task completion.", "Users complete critical tasks without undocumented assistance."),
        ("Sales demo validation", "Run the approved demo story on a clean demo tenant.", "Demo data, speed and screens are stable and presentable."),
        ("Known-gap acceptance", "Review every pending/gap-tagged case.", "Each gap is fixed, excluded, or formally accepted with owner/date."),
        ("Go-live checklist", "Review domains, backups, monitoring, support, data, users and rollback.", "All mandatory checklist items have evidence and owner sign-off."),
        ("Regression completion", "Execute the complete release regression set.", "Required execution/pass thresholds are met."),
        ("Final release decision", "Review dashboard, blockers, risks and sign-offs.", "READY FOR SALES appears only when every configured gate is met."),
    ], default_type="UAT", default_priority="Critical")


def level_intro(level):
    if level == 1:
        return (
            "Core Flow and Smoke Testing",
            "Proves that the main LIMS workflow works from lab setup to patient and doctor report access.",
            ["A stable QA build is deployed.", "Required roles and master data can be created.", "No known blocker prevents the main workflow."],
            ["100% of Level 1 cases executed.", "100% pass rate for applicable cases.", "No open Critical/High failure."]
        )
    if level == 2:
        return (
            "Functional, Integration, Negative and Gap Testing",
            "Tests detailed rules, boundaries, errors, integrations, permissions, and every important gap missing from the original flow.",
            ["Level 1 has passed.", "Feature-complete QA build and representative data are available.", "Product decisions are documented for ambiguous rules."],
            ["At least 95% execution and 95% pass for applicable cases.", "All Critical/High failures are closed or formally rejected with evidence.", "Every gap-tagged case has a decision, owner and target release."],
        )
    return (
        "Market Readiness and Release Certification",
        "Acts as the sales-release gate. It proves security, privacy, scale, recovery, compatibility, integrity, production operations, and real-user readiness.",
        ["Levels 1 and 2 meet their exit criteria.", "Release candidate is deployed in a production-like environment.", "Backup, monitoring, support and rollback owners are available."],
        ["100% of Level 3 cases executed.", "At least 98% pass for applicable Level 3 cases.", "100% evidence for release-gate cases.", "Zero open Critical/High failure.", "Final dashboard status is READY FOR SALES."],
    )


def build_level_doc(level):
    title, purpose, entry, exit_rules = level_intro(level)
    doc = Document()
    configure(doc, f"LIMS | LEVEL {level} TESTING", f"LIMS Level {level} - {title}")
    title_page(doc, f"LEVEL {level} TESTING", f"LIMS Level {level}: {title}", purpose,
               "Manual QA, UAT and release teams. Use the consolidated Excel workbook to record execution results.")
    add_heading(doc, "1. Objective", 1)
    add_body(doc, purpose)
    add_heading(doc, "2. Entry criteria", 1)
    bullets(doc, entry)
    add_heading(doc, "3. Exit criteria", 1)
    bullets(doc, exit_rules)
    if level == 1:
        add_heading(doc, "4. Missed flow steps now included", 1)
        bullets(doc, [
            "Lab tenant, plan and module setup before lab work starts.",
            "Role separation for collection, result entry and report release.",
            "Test master, reference ranges, package and inventory setup.",
            "Partial/final payment, invoice, receipt, ledger and commission checks.",
            "Barcode, two-identifier collection, custody, inventory reservation and status synchronization.",
            "Report review, approval, release, notification and released-only portal access.",
            "Patient/doctor activation, ownership, tenant isolation and audit evidence.",
        ])
    else:
        add_heading(doc, "4. Testing approach", 1)
        numbered(doc, [
            "Use the Case ID from this document and the Excel tracker.",
            "Prepare the stated role, data and preconditions.",
            "Follow the action steps and save evidence.",
            "Record Passed, Failed, Blocked, Not Run or N/A in the workbook.",
            "Create a defect for every unexpected result.",
            "Retest fixes and update defect status before release review.",
        ])
    add_heading(doc, "5. Scenario catalogue", 1)
    level_cases = [c for c in CASES if c["Level"] == f"Level {level}"]
    by_module = defaultdict(list)
    for case in level_cases:
        by_module[case["Module"]].append(case)
    for module, items in by_module.items():
        add_heading(doc, module, 2)
        rows = []
        for c in items:
            expected = c["Expected Result"]
            if c["Gap Covered"]:
                expected += f" Gap/decision: {c['Gap Covered']}."
            rows.append((c["Case ID"], c["Scenario"], c["Steps"], expected))
        add_table(doc, ["ID", "Scenario", "Steps", "Pass condition"], rows,
                  [850, 2000, 3250, 3260], font_size=7.7 if level > 1 else 8.1)
    add_heading(doc, "6. Result recording", 1)
    add_body(doc, "Record all results in the consolidated Excel workbook. Do not calculate readiness manually. The workbook uses formulas, release thresholds and blocker rules to decide whether the product is ready for sales.")
    doc.save(DOCS[level])


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    QA.mkdir(parents=True, exist_ok=True)
    build_catalog()
    CATALOG.write_text(json.dumps(CASES, indent=2), encoding="utf-8")
    for level in (1, 2, 3):
        build_level_doc(level)
    print(json.dumps({"cases": len(CASES), "by_level": dict(COUNTERS), "catalog": str(CATALOG), "docs": {k: str(v) for k, v in DOCS.items()}}, indent=2))


if __name__ == "__main__":
    main()
